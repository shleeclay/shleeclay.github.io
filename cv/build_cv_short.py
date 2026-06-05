# -*- coding: utf-8 -*-
"""
Two-page condensed CV — Seunghyeon (Clay) Lee.
Publications + manuscripts kept in full (APA, IF/Q, clickable DOI); everything else one-line.
Run:  python build_cv_short.py  ->  Lee_Seunghyeon_CV_short.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, json

HERE = os.path.dirname(os.path.abspath(__file__))
INK = RGBColor(0x1a, 0x1a, 0x1a); ACCENT = RGBColor(0x1f, 0x4e, 0x3d)
GREY = RGBColor(0x55, 0x55, 0x55); LIGHT = RGBColor(0x88, 0x88, 0x88)
F = "Calibri"

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.45); s.bottom_margin = Inches(0.45)
    s.left_margin = Inches(0.6); s.right_margin = Inches(0.6)
RIGHT_TAB = Inches(7.3)
st = doc.styles["Normal"]; st.font.name = F; st.font.size = Pt(8.8); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.0

def run(p, t, size=8.8, bold=False, italic=False, color=INK, caps=False, spacing=None):
    r = p.add_run(t); r.font.name = F; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr(); sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing)); rPr.append(sp)
    return r

def para(sb=0, sa=0, indent=None, hang=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
        if hang: p.paragraph_format.first_line_indent = Inches(-indent)
    return p

def rtab(p): p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

def bborder(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1f4e3d')
    pbdr.append(b); pPr.append(pbdr)

def hyperlink(p, url, text, color="555555", size=8.3):
    rid = p.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), rid)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr'); rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), F)
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); link.append(r); p._p.append(link)

def section(title):
    p = para(6, 2); run(p, title, 9.5, bold=True, color=ACCENT, caps=True, spacing=20); bborder(p)

# ----------------------------------------------------------------- header
p = para(0, 1); run(p, "Seunghyeon Lee", 16, bold=True); run(p, "  (이승현 · Clay Lee)", 10, color=GREY)
p = para(0, 2)
run(p, "Ph.D. Candidate (expected Aug 2026) · Forest Structure & Remote Sensing (LiDAR / GEDI) · "
       "Seoul National University", 9, bold=True, color=ACCENT)
contact = [
    ("shlee5598@snu.ac.kr", "mailto:shlee5598@snu.ac.kr"), ("+82-10-8455-5598", None),
    ("drseunghyeonlee.com", "https://drseunghyeonlee.com/en"),
    ("ORCID 0000-0003-4612-065X", "https://orcid.org/0000-0003-4612-065X"),
    ("Google Scholar", "https://scholar.google.com/citations?user=Ew1_-r0AAAAJ"),
]
p = para(0, 4)
for i, (disp, url) in enumerate(contact):
    if i: run(p, "   |   ", 8.5, color=LIGHT)
    hyperlink(p, url, disp, size=8.5) if url else run(p, disp, 8.5, color=GREY)
bborder(p)

# ----------------------------------------------------------------- education
section("Education")
edu = [
    ("Ph.D. Candidate, Interdisciplinary Program in Landscape Architecture, Seoul National University",
     "2022.03 — Aug 2026 (expected)",
     "Dissertation: “LiDAR-Based Characterization of Vertical Structure and Disturbance in Temperate "
     "Forests” · Advisor: Prof. Youngkeun Song"),
    ("M.A., Landscape Architecture, Seoul National University", "2020.03 — 2022.02", None),
    ("B.S., Landscape Architecture, Kyungpook National University", "2011.03 — 2017.02", None),
]
for lead, date, sub in edu:
    p = para(2, 0); rtab(p); run(p, lead, 8.8, bold=True); run(p, "\t" + date, 8.5, color=GREY)
    if sub:
        p = para(0, 0); run(p, sub, 8.3, italic=True, color=GREY)

# ----------------------------------------------------------------- research experience
section("Research Experience")
for lead, date, sub in [
    ("Visiting Research Intern — Purdue University, FNR FUSE Lab (PI: Prof. Brady Hardiman)",
     "2023.12 — 2024.02", "West Lafayette, IN, USA · forest structure & LiDAR"),
    ("Visiting Research Intern — Virginia Tech (PI: Prof. Jaeyoung Ha)",
     "2024.07 — 2024.08", "Blacksburg, VA, USA · landscape architecture")]:
    p = para(2, 0); rtab(p); run(p, lead, 8.8, bold=True); run(p, "\t" + date, 8.5, color=GREY)
    p = para(0, 0); run(p, sub, 8.3, italic=True, color=GREY)

# ----------------------------------------------------------------- publications (full, APA)
def _bib_entries(path):
    txt = open(path, encoding="utf-8").read(); out, i = [], 0
    while True:
        at = txt.find("@", i)
        if at < 0: break
        b = txt.find("{", at); depth, j = 0, b
        while j < len(txt):
            if txt[j] == "{": depth += 1
            elif txt[j] == "}":
                depth -= 1
                if depth == 0: break
            j += 1
        out.append(txt[at:j + 1]); i = j + 1
    return out

def _bf(raw, name):
    m = re.search(r"\b" + name + r"\s*=\s*\{(.*?)\}\s*[,}\n]", raw, re.I | re.S)
    return m.group(1).strip() if m else ""

def _initials(g):
    return " ".join("-".join(s[0].upper() + "." for s in part.split("-") if s) for part in g.split())

def _apa(name):
    if "," in name:
        sur, giv = [x.strip() for x in name.split(",", 1)]
    else:
        ps = name.split(); sur, giv = ps[-1], " ".join(ps[:-1])
    return f"{sur}, {_initials(giv)}".strip().rstrip(","), (sur == "Lee" and giv.startswith("Seunghyeon"))

_MO = {m: i for i, m in enumerate(["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"], 1)}
def _mon(raw):
    m = _bf(raw, "month").lower()[:3]
    return _MO.get(m, int(m) if m.isdigit() else 0)

_site = json.load(open(os.path.join(HERE, "..", "src", "data", "site.json"), encoding="utf-8"))
IFQ = {(it.get("doi") or "").lower().strip(): (it.get("if", ""), it.get("quartile", ""))
       for it in _site["publications"]["items"] if it.get("doi")}

PS = 8.3   # publication font size
def render_pub(n, apa, year, title, journal, vol, num, pages, doi, iff, q):
    p = para(2, 0, indent=0.22, hang=True)
    run(p, f"{n}. ", PS, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i: run(p, ", ", PS)
        if i == nA - 1 and nA > 1: run(p, "& ", PS)
        run(p, disp, PS, bold=is_user)
        if i == 0 and is_user: run(p, "†", PS, bold=True, color=ACCENT)
    run(p, f" ({year}). ", PS); run(p, title + ". ", PS); run(p, journal, PS, italic=True)
    if vol:
        run(p, ", ", PS); run(p, vol, PS, italic=True)
        if num: run(p, f"({num})", PS)
        if pages: run(p, f", {pages}", PS)
    run(p, ". ", PS)
    if doi: hyperlink(p, f"https://doi.org/{doi}", f"https://doi.org/{doi}", size=PS)
    if iff or q:
        iv = iff.split("(")[0].strip()
        run(p, "  · " + ", ".join(x for x in [f"IF {iv}" if iv else "", q] if x), 7.8, color=ACCENT)

section("Peer-Reviewed Publications")
p = para(0, 2); run(p, "APA 7th · name in bold = author · † = first author · SCI then KCI, newest first.", 7.8, italic=True, color=LIGHT)
_pubs = []
for raw in _bib_entries(os.path.join(HERE, "citations", "publications.bib")):
    doi = _bf(raw, "doi"); iff, q = IFQ.get(doi.lower(), ("", ""))
    yv = _bf(raw, "year"); yv = int(yv) if yv.isdigit() else 0
    _pg = re.search(r"\d+", _bf(raw, "pages")); sp = int(_pg.group()) if _pg else 0
    _pubs.append({
        "apa": [_apa(a.strip()) for a in _bf(raw, "author").split(" and ")], "year": _bf(raw, "year"),
        "title": re.sub(r"\s*-\s*$", "", _bf(raw, "title")).replace(" - ", ": ").rstrip(". "),
        "journal": _bf(raw, "journal").replace("\\&", "&"), "vol": _bf(raw, "volume"),
        "num": _bf(raw, "number"), "pages": _bf(raw, "pages").replace("-", "–"), "doi": doi,
        "iff": iff, "q": q, "sortk": (0 if (iff or q) else 1, -yv, -_mon(raw), sp),
    })
_pubs.sort(key=lambda d: d["sortk"])
for n, d in enumerate(_pubs, 1):
    render_pub(n, d["apa"], d["year"], d["title"], d["journal"], d["vol"], d["num"], d["pages"], d["doi"], d["iff"], d["q"])

# ----------------------------------------------------------------- manuscripts under review
section("Manuscripts Under Review")
under_review = [
    (["Seunghyeon Lee", "Dennis Heejoon Choi", "Youngkeun Song", "James H. Thorne"],
     "Spaceborne LiDAR reveals post-fire vertical structural loss in a dense temperate Asian conifer plantation.",
     "Science of Remote Sensing", "Manuscript under review"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"],
     "Systematic bias in urban-forest vegetation coverage assessment: the influence of topo-edaphic conditions on discrepancies between ALS and field surveys.",
     "Geoscience Letters", "Revise and resubmit, 1st round"),
    (["Seunghyeon Lee", "Yonghwan Kim", "Dohee Kim", "Hansoo Kim", "Youngkeun Song"],
     "Bi-temporal ALS assessment of vertical–horizontal forest structures and their structural association in a temperate urban forest.",
     "Forest Science and Technology", "Revise and resubmit, 2nd round"),
    (["Yonghwan Kim", "Seunghyeon Lee", "Wonhyeop Shin", "Youngkeun Song"],
     "Integrating UAV-derived habitat metrics with movement persistence and species distribution models to characterize seasonal habitat use of invasive turtles in urban wetlands.",
     "Global Ecology and Conservation", "Revise and resubmit, 1st round"),
    (["Gapseong Jekal", "Yong Hwan Kim", "Seunghyeon Lee", "Ji Weon Yun", "Dae Yeol Kim", "Youngkeun Song"],
     "Monitoring transition-zone dynamics of a Phragmites communis–Suaeda japonica mosaic from multi-season Sentinel-2 in a coastal wetland.",
     "Journal of Coastal Conservation", "Revise and resubmit, 1st round"),
]
for n, (auth, title, jour, stage) in enumerate(under_review, 1):
    apa = [_apa(a) for a in auth]
    p = para(2, 0, indent=0.22, hang=True); run(p, f"{n}. ", PS, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i: run(p, ", ", PS)
        if i == nA - 1 and nA > 1: run(p, "& ", PS)
        run(p, disp, PS, bold=is_user)
        if i == 0 and is_user: run(p, "†", PS, bold=True, color=ACCENT)
    run(p, " (2026). ", PS); run(p, title.rstrip(". "), PS, italic=True)
    run(p, f" [{stage}]. ", PS); run(p, jour + ".", PS)

# ----------------------------------------------------------------- patents
section("Patents (Republic of Korea)")
p = para(2, 0); run(p, "11 patents — 7 registered, 4 pending.  ", 8.5, bold=True)
run(p, "Highlights: drone-LiDAR wetland vegetation-structure determination (reg. 2025); spaceborne-LiDAR "
       "wildfire damage assessment (pending); tree-management platform (reg. 2025, sole inventor).", 8.5, color=GREY)

# ----------------------------------------------------------------- projects (selected)
section("Research Projects (selected)")
projs = [
    ("Integrated assessment tool for carbon reduction & ecosystem services", "Korean Ministry of Environment", "2023 — 2027", "Researcher", "₩1,579M"),
    ("Real-time positioning surveillance system for invasive species", "Korean Ministry of Environment", "2021 — 2023", "Project Manager", "₩1,208M"),
    ("Redundancy-based green-infrastructure technology for urban ecosystems", "Korean Ministry of Environment", "2020 — 2022", "Researcher", "₩927M"),
    ("Carbon storage/sink analysis & inventory for Gyeonggi-do by biotope", "Gyeonggi Research Institute", "2024 — 2025", "Researcher", None),
    ("Gwacheon-si urban ecological status mapping", "Gyeonggi Research Institute", "2021 — 2022", "Researcher", None),
]
for title, funder, period, role, budget in projs:
    p = para(2, 0); rtab(p); run(p, title, 8.5, bold=True); run(p, "\t" + period, 8.3, color=GREY)
    sub = f"{funder} · {role}" + (f" · {budget} (total)" if budget else "")
    p = para(0, 0); run(p, sub, 8.3, italic=True, color=GREY)
p = para(1, 0); run(p, "…and 7 additional national / municipal R&D projects (2018 — present).", 8.0, italic=True, color=LIGHT)

# ----------------------------------------------------------------- conferences + invited (summary)
section("Conference Presentations & Invited Talks")
p = para(2, 0)
run(p, "International: ", 8.5, bold=True)
run(p, "13 presentations (AGU ×4, ESA ×2, ICLEE ×2, World Forestry Congress, JCK) + 10 domestic "
       "(KOSERT / KSEE / KILA).  ", 8.5)
run(p, "Invited: NEF (Orlando, 2024), Asia Week (Fukuoka, 2024), SNU guest lectures (2025).  "
       "Full list on request.", 8.5, color=GREY)

# ----------------------------------------------------------------- awards
section("Awards & Honors")
p = para(2, 0); run(p, "Best Presentation Award, ", 8.5, bold=True); run(p, "KOSERT (2022, 2025).  ", 8.5)
run(p, "Full graduate scholarships: Chung Mong-Koo Foundation (2022–), BK21 FOUR (2020–); "
       "Ilju Foundation undergraduate scholarship (2015–17).", 8.5, color=GREY)

# ----------------------------------------------------------------- teaching / service / professional
section("Teaching · Service · Professional")
p = para(2, 0); run(p, "Teaching:  ", 8.5, bold=True, color=ACCENT)
run(p, "University Lecturer, Incheon National University (2024–25); 4 online GIS/QGIS courses "
       "(1,000+ students); 20+ invited workshops.", 8.5)
p = para(1, 0); run(p, "Service:  ", 8.5, bold=True, color=ACCENT)
run(p, "Peer reviewer, Int. J. Applied Earth Observation & Geoinformation.  Member: AGU, ESA, KOSERT, "
       "Ecological Society of Korea.", 8.5)
p = para(1, 0); run(p, "Professional:  ", 8.5, bold=True, color=ACCENT)
run(p, "Founder, TREE:ID — urban tree-inventory & management SaaS startup (2023–25); "
       "Landscape Architect, SUNJIN E&A (2017–18).", 8.5)

# ----------------------------------------------------------------- skills
section("Technical Skills")
for label, val in [
    ("Programming", "Python, R, Google Earth Engine"),
    ("LiDAR / RS", "ALS · GEDI · MLS, LAStools / PDAL, canopy & vertical-structure metrics; QGIS / ArcGIS, "
                   "UAV & thermal imaging, ENVI-met"),
    ("Machine learning", "XGBoost, deep learning for forest & wildlife mapping"),
    ("Languages", "Korean (native), English (professional working)"),
    ("Certifications", "Engineer Forest, Engineer Landscape Architecture, Drone Pilot (Class 2)"),
]:
    p = para(1, 0); run(p, label + ":  ", 8.5, bold=True, color=ACCENT); run(p, val, 8.5)

out = "Lee_Seunghyeon_CV_short.docx"
doc.save(out)
print("saved:", out)
