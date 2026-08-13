# -*- coding: utf-8 -*-
"""CV 회귀 검증 — build_cv.py / build_cv_short.py 의 출력 텍스트를 기준선과 대조한다.

리팩터 중 문장이 빠지거나 순서가 바뀌는 것을 잡기 위한 장치.
빌드 스크립트는 **임시 폴더에서 실행**하므로 저장소의 docx 는 건드리지 않는다.
생성 날짜처럼 매번 달라지는 줄은 정규화해서 오탐을 막는다.

  python _cv_baseline.py --save     기준선 저장 (리팩터 전에 1회)
  python _cv_baseline.py --check    현재 출력과 기준선 대조
"""
import argparse
import difflib
import os
import re
import shutil
import subprocess
import sys
import tempfile

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
BASE = os.path.join(HERE, "_baseline")

TARGETS = [
    ("build_cv.py", "Lee_Seunghyeon_CV_v9.docx", "cv_full.txt"),
    ("build_cv_short.py", "Lee_Seunghyeon_CV_short.docx", "cv_short.txt"),
]

# 매 실행마다 달라지는 부분 — 대조에서 제외
VOLATILE = [
    re.compile(r"Date:\s*\d{1,2}\s+\w+\s+\d{4}"),
    re.compile(r"Signature:\s*_+"),
]


def build(script, outname):
    """임시 폴더에서 빌드 스크립트를 돌리고 생성된 docx 경로를 반환."""
    tmp = tempfile.mkdtemp(prefix="cvbase_")
    r = subprocess.run([sys.executable, os.path.join(HERE, script)],
                       cwd=tmp, capture_output=True, text=True, encoding="utf-8")
    if r.returncode != 0:
        shutil.rmtree(tmp, ignore_errors=True)
        raise SystemExit(f"[{script}] 빌드 실패\n{r.stdout}\n{r.stderr}")
    path = os.path.join(tmp, outname)
    if not os.path.exists(path):
        shutil.rmtree(tmp, ignore_errors=True)
        raise SystemExit(f"[{script}] 산출물 없음: {outname}")
    return tmp, path


def extract(path):
    """문단 텍스트를 줄 단위로. 표가 있으면 셀도 포함."""
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    out = []
    for l in lines:
        for pat in VOLATILE:
            l = pat.sub("<VOLATILE>", l)
        out.append(l)
    return out


def run(mode):
    os.makedirs(BASE, exist_ok=True)
    bad = 0
    for script, outname, basename in TARGETS:
        if not os.path.exists(os.path.join(HERE, script)):
            print(f"  {script}: 없음 — 건너뜀")
            continue
        tmp, docx = build(script, outname)
        try:
            cur = extract(docx)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
        ref = os.path.join(BASE, basename)

        if mode == "save":
            with open(ref, "w", encoding="utf-8", newline="\n") as f:
                f.write("\n".join(cur) + "\n")
            print(f"  {script:22s} -> _baseline/{basename}  ({len(cur)}줄)")
        else:
            if not os.path.exists(ref):
                print(f"  {script:22s} 기준선 없음 — 먼저 --save 실행")
                bad += 1
                continue
            old = open(ref, encoding="utf-8").read().splitlines()
            d = list(difflib.unified_diff(old, cur, "baseline", "current", lineterm="", n=1))
            if d:
                bad += 1
                print(f"\n  {script} — 차이 {sum(1 for l in d if l[:1] in '+-' and l[:3] not in ('+++','---'))}줄")
                for l in d:
                    print("   ", l[:150])
            else:
                print(f"  {script:22s} 기준선과 동일 ({len(cur)}줄)")
    return bad


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--save", action="store_true")
    g.add_argument("--check", action="store_true")
    a = ap.parse_args()
    sys.exit(run("save" if a.save else "check"))
