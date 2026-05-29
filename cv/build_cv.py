# -*- coding: utf-8 -*-
"""
Academic CV builder — Seunghyeon (Clay) Lee
Forest remote sensing / LiDAR · GEDI / ecology — postdoc application

Structure follows field-standard academic CVs (Pascual UMD / Harvard GSAS).
All content is sourced from the verified portfolio data in
  ../src/data/site.json
Edit the data blocks below, then run:  python build_cv.py
Output: Lee_Seunghyeon_CV_v1.docx  ->  export to PDF in Word for ../public/cv/
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- palette
INK    = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x1f, 0x4e, 0x3d)   # deep forest green
GREY   = RGBColor(0x55, 0x55, 0x55)
LIGHT  = RGBColor(0x88, 0x88, 0x88)
BODYFONT = "Calibri"
HEADFONT = "Calibri"

doc = Document()

# page margins
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

RIGHT_TAB = Inches(6.9)  # within 7.0" text width

# base style
st = doc.styles["Normal"]
st.font.name = BODYFONT
st.font.size = Pt(9.5)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.04

def _set_run(r, size=9.5, bold=False, italic=False, color=INK, font=BODYFONT, caps=False, spacing=None):
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr(); sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing)); rPr.append(sp)

def add_right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

def bottom_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1f4e3d')
    pbdr.append(b); pPr.append(pbdr)

# ---------------------------------------------------------------- header
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Seunghyeon Lee"); _set_run(r, 22, bold=True, color=INK)
r = p.add_run("  (이승현 · Clay Lee)"); _set_run(r, 12, color=GREY)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Ph.D. Candidate (expected Aug 2026) · Forest Structure & Remote Sensing (LiDAR / GEDI)")
_set_run(r, 10, color=ACCENT, bold=True)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Landscape & Ecological Planning Lab, Seoul National University, Seoul, Republic of Korea")
_set_run(r, 9.5, color=GREY)
p.paragraph_format.space_after = Pt(2)

contact = [
    "shlee5598@snu.ac.kr",
    "+82-10-8455-5598",
    "shleeclay.github.io",
    "github.com/shleeclay",
    "Google Scholar / ORCID: [add URL]",
]
p = doc.add_paragraph()
for i, c in enumerate(contact):
    if i:
        sep = p.add_run("   |   "); _set_run(sep, 9, color=LIGHT)
    r = p.add_run(c); _set_run(r, 9, color=GREY)
bottom_border(p); p.paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------------- section helper
def section(title):
    p = doc.add_paragraph()
    r = p.add_run(title); _set_run(r, 11, bold=True, color=ACCENT, caps=True, spacing=30)
    bottom_border(p)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)

def entry(left_bold, right="", sub="", bullets=None, left_rest=""):
    """One CV entry: bold lead + optional right-aligned date, italic subline, bullets."""
    p = doc.add_paragraph(); add_right_tab(p)
    r = p.add_run(left_bold); _set_run(r, 9.5, bold=True)
    if left_rest:
        r = p.add_run(left_rest); _set_run(r, 9.5)
    if right:
        r = p.add_run("\t" + right); _set_run(r, 9, color=GREY)
    p.paragraph_format.space_before = Pt(3)
    if sub:
        ps = doc.add_paragraph(); r = ps.add_run(sub); _set_run(r, 9, italic=True, color=GREY)
    if bullets:
        for b in bullets:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent = Inches(0.25)
            pb.paragraph_format.space_after = Pt(0)
            r = pb.add_run(b); _set_run(r, 9)

def authors_run(p, authors, me="Seunghyeon Lee", mark_first=False):
    """Render byline; bold `me`. If mark_first and `me` is the first author, add † after the name."""
    for i, a in enumerate(authors):
        if i:
            s = p.add_run(", "); _set_run(s, 9, color=INK)
        r = p.add_run(a); _set_run(r, 9, bold=(a == me))
        if mark_first and i == 0 and a == me:
            d = p.add_run("†"); _set_run(d, 9, bold=True, color=ACCENT)

# ---------------------------------------------------------------- Research Interests
section("Research Interests")
p = doc.add_paragraph()
r = p.add_run("Airborne & spaceborne LiDAR (GEDI) for vertical forest structure · forest typology and "
              "stratification metrics (FHD, PAI, layer indices) · biomass and carbon estimation · "
              "urban ecosystems and green infrastructure · UAV/thermal wildlife monitoring · "
              "reproducible, code-driven remote-sensing science.")
_set_run(r, 9.5)
p.paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- Education
section("Education")
edu = [
    ("Ph.D. Candidate, Interdisciplinary Program in Landscape Architecture",
     "2022.03 — Aug 2026 (expected)",
     "Integrated Major in SmartCity Global Convergence · Seoul National University · "
     "Landscape & Ecological Planning Lab · Advisor: Prof. Youngkeun Song",
     ["Dissertation: [title — to be added]"]),
    ("M.A., Landscape Architecture",
     "2020.03 — 2022.02",
     "Integrated Major in SmartCity Global Convergence · Graduate School of Environmental Studies, "
     "Seoul National University · Advisor: Prof. Youngkeun Song", None),
    ("B.S., Landscape Architecture",
     "2011.03 — 2017.02",
     "College of Agriculture & Life Sciences, Kyungpook National University, Daegu, Republic of Korea", None),
]
for lead, date, sub, bl in edu:
    entry(lead, right=date, sub=sub, bullets=bl)

# ---------------------------------------------------------------- Research Experience
section("Research Experience")
exp = [
    ("Visiting Research Intern", " — Purdue University, FNR FUSE Lab (PI: Prof. Brady Hardiman)",
     "2023.12 — 2024.02", "West Lafayette, IN, United States · Forest structure & LiDAR"),
    ("Visiting Research Intern", " — Virginia Tech (PI: Prof. Jaeyoung Ha)",
     "2024.07 — 2024.08", "Blacksburg, VA, United States · Landscape architecture"),
    ("Graduate Researcher", " — Landscape & Ecological Planning Lab, Seoul National University",
     "2020.03 — Present", "National R&D and municipal projects on LiDAR forest structure, carbon, and urban ecology"),
    ("Research Assistant", " — Architecture & Urban Research Institute (AURI)",
     "2019.08 — 2019.11", "Republic of Korea"),
]
for lead, rest, date, sub in exp:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- Publications
section("Peer-Reviewed Publications")
p = doc.add_paragraph()
r = p.add_run("Name in bold = author.  † = first author.  (9 publications; reverse chronological)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

pubs = [
    (True, ["Gukhwa Jang", "Seunghyeon Lee", "Sung-Ho Kil", "Youngkeun Song"], 2026,
     "Comparing LST estimation methods under matched spatiotemporal conditions in low- and high-rise "
     "residential areas: satellite, simulation, and UAV approaches.", "KOSERT", ""),
    (True, ["Seunghyeon Lee", "Youngkeun Song"], 2026,
     "Forest-type and seasonal analysis of GEDI–ALS canopy height agreement and GEDI structural metrics "
     "(FHD, PAI): a case study of forests in Gwacheon-si and Uiwang-si.", "KOSERT", ""),
    (True, ["Seunghyeon Lee", "Uirin Ha", "Heejae Lee", "Hyeyeong Choe", "Hansoo Kim", "Youngkeun Song"], 2026,
     "Multi-scale forest typologies using novel vertical layer metrics from airborne LiDAR in temperate "
     "mixed forests.", "Ecological Indicators", "IF 7.4, Q1 · doi:10.1016/j.ecolind.2026.114798"),
    (False, ["Gapseong Jekal", "Seunghyeon Lee", "Youngeun Yang", "Youngkeun Song"], 2026,
     "Non-destructive carbon storage estimation of Salix spp. community in wetlands using LiDAR: a case "
     "study of Ungok Wetland.", "KOSERT", "doi:10.13087/kosert.2025.28.6.95"),
    (False, ["Jinyuan Shao", "Dennis Heejoon Choi", "Jidong Liu", "Xiangxi Tian", "Bina Thapa",
             "Seunghyeon Lee", "Ayman Habib", "Songlin Fei"], 2026,
     "A three-stage framework for stand-level automated stem volume estimation in temperate forests using "
     "mobile laser scanning.", "Remote Sensing of Environment", "IF 11.4, Q1 · doi:10.1016/j.rse.2026.115246"),
    (False, ["Jaeyeon Kim", "Seungwoo Han", "Jiweon Yun", "Seunghyeon Lee", "Youngkeun Song"], 2026,
     "Ecological structures and terrestrial insect diversity across successional stages in abandoned paddy "
     "fields.", "Agriculture, Ecosystems & Environment", "IF 6.4, Q1 · doi:10.1016/j.agee.2025.110172"),
    (False, ["Jiweon Yun", "Seunghyeon Lee", "Youngkeun Song"], 2025,
     "Assessing Corvus frugilegus (Rook) habitat preferences through flock-size-specific species "
     "distribution modeling using citizen science data.", "Global Ecology and Conservation",
     "IF 3.4, Q1 · doi:10.1016/j.gecco.2025.e03866"),
    (False, ["Younha Han", "Wonhyeop Shin", "Jihwan Kim", "Dohee Kim", "Jiweon Yun", "Sokyoung Yi",
             "Yonghwan Kim", "Seunghyeon Lee", "Youngkeun Song"], 2024,
     "Diel activity patterns of water deer (Hydropotes inermis) and wild boar (Sus scrofa) in a suburban "
     "area monitored by long-term camera-trapping.", "KOSERT", ""),
    (True, ["Seunghyeon Lee", "Youngkeun Song", "Sung-Ho Kil"], 2021,
     "Feasibility analyses of real-time detection of wildlife using UAV-derived thermal and RGB images.",
     "Remote Sensing", "IF 5.0, Q1 · doi:10.3390/rs13112169"),
]
for n, (first, auth, yr, title, jour, meta) in enumerate(pubs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    lead = p.add_run(f"{len(pubs)-n+1}. "); _set_run(lead, 9, color=GREY)
    authors_run(p, auth, mark_first=True)
    yr_r = p.add_run(f" ({yr}). "); _set_run(yr_r, 9)
    t = p.add_run(title + " "); _set_run(t, 9)
    j = p.add_run(jour + ".") ; _set_run(j, 9, italic=True)
    if meta:
        m = p.add_run("  " + meta); _set_run(m, 8.5, color=LIGHT)

# ---------------------------------------------------------------- Conferences
section("Selected Conference Presentations")
p = doc.add_paragraph()
r = p.add_run("International first-author / presenter highlights (full list of 25 available on request).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

confs = [
    (["Hansoo Kim", "Seunghyeon Lee", "Youngkeun Song"], 2025, "Poster",
     "Novel LiDAR indices reveal scale-dependent vertical structure and typologies in a temperate forest.",
     "AGU Fall Meeting, New Orleans, US"),
    (["Hansoo Kim", "Seunghyeon Lee", "Youngkeun Song"], 2024, "Poster",
     "Quantifying the number of forest stratification layers of city-scale forest and characterizing by "
     "forest types.", "AGU Fall Meeting, Washington D.C., US"),
    (["Hansoo Kim", "Seunghyeon Lee", "Youngkeun Song"], 2024, "Oral",
     "Canopy layers distribution by forest patch and species at the city-level forest using airborne LiDAR.",
     "ICLEE, Kitakyushu, Japan"),
    (["Hansoo Kim", "Seunghyeon Lee", "Youngkeun Song"], 2024, "Poster",
     "Characterizing forest vegetation stratification by forest biotope types using airborne LiDAR.",
     "ESA Annual Meeting, Long Beach, CA, US"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2023, "Poster",
     "The impacts of fire-induced disturbances on tree height and structure using GEDI-derived variables.",
     "ESA Annual Meeting, Portland, OR, US"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Comparison of GEDI and aerial laser scanning datasets according to leaf-on and leaf-off seasons.",
     "AGU Fall Meeting, Chicago, IL, US / Online"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Monitoring structural change of fire-induced forest vegetation using GEDI.", "ICLEE (12th), Online"),
    (["Seunghyeon Lee", "Sung-Ho Kil", "Youngkeun Song"], 2022, "Poster",
     "Feasibility analyses of real-time detection of wildlife using UAV-derived thermal and RGB images.",
     "World Forestry Congress (15th), Seoul, Republic of Korea"),
]
for auth, yr, kind, title, venue in confs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
    authors_run(p, auth)
    yr_r = p.add_run(f" ({yr}). "); _set_run(yr_r, 9)
    t = p.add_run(title + " "); _set_run(t, 9)
    v = p.add_run(f"[{kind}] {venue}."); _set_run(v, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Patents
section("Patents")
p = doc.add_paragraph()
r = p.add_run("Republic of Korea — 6 registered, 5 pending (selected; full list available on request).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
patents = [
    ("A method for determining wetland vegetation structure using drone LiDAR and a system for "
     "implementing it.", "Registered 2025.04 · Seoul National University"),
    ("Method and system for quantitative assessment and automatic classification of wildfire damage using "
     "spaceborne LiDAR.", "Pending (2025) · Dankook Univ. & Seoul National University"),
    ("System and method for providing a tree management platform.", "Registered 2025.09 · Sole inventor"),
    ("Image-based roadside-tree information management system and method.", "Registered 2024.11 · Sole inventor"),
    ("Automated wild-animal detection method and device using drone-mounted thermal camera imagery.",
     "Registered 2022.05 · Seoul National University"),
]
for title, meta in patents:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
    t = p.add_run(title + " "); _set_run(t, 9)
    m = p.add_run(meta); _set_run(m, 8.5, italic=True, color=GREY)

# ---------------------------------------------------------------- Grants / Projects
section("Research Projects & Grants")
projects = [
    ("Integrated assessment tool for carbon reduction by offset and synergy with ecosystem services",
     "Korean Ministry of Environment", "2023 — 2027", "Researcher"),
    ("Creation, restoration & management technology of carbon-accumulated abandoned-paddy wetland",
     "Korean Ministry of Environment", "2022 — 2026", "Researcher"),
    ("Carbon storage/sink analysis and inventory for Gyeonggi-do considering biotope types",
     "Gyeonggi Research Institute", "2024 — 2025", "Researcher"),
    ("Real-time web-based positioning surveillance system for introduced exotic species",
     "Korean Ministry of Environment", "2021 — 2023", "Project Manager"),
    ("Redundancy-based green-infrastructure technology for urban ecosystem challenges",
     "Korean Ministry of Environment", "2020 — 2022", "Researcher"),
]
for title, inst, period, role in projects:
    entry(title, right=period, sub=f"{inst} · {role}")
p = doc.add_paragraph(); r = p.add_run("…and 7 additional national/municipal R&D projects (2018 — present).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_before = Pt(2)

# ---------------------------------------------------------------- Honors
section("Honors & Scholarships")
honors = [
    ("Chung Mong-Koo Foundation — OnDream Future Industry Talent Graduate Scholarship (Full)", "2022 — Present"),
    ("BK21 FOUR — Integrated Major in SmartCity Global Convergence (Full)", "2020 — Present"),
    ("Ilju Foundation Undergraduate Scholarship, 23rd Scholar (Full)", "2015 — 2017"),
    ("Challenge Scholarship, Kyungpook National University (Full)", "2011 — 2013"),
]
for h, d in honors:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(h); _set_run(r, 9)
    r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

# ---------------------------------------------------------------- Teaching
section("Teaching & Outreach")
teach = [
    ("University Lecturer — “Understanding and Application of Spatial Information” (3 cr.), "
     "Incheon National University", "2024, 2025"),
    ("Online Instructor — 4 GIS/QGIS courses on Inflearn (750+ enrolled students total)", "2022 — 2025"),
    ("Invited Workshops — K-water, Korea Environment Corporation, Hyundai NGV, SeSAC, SNU (20+ sessions)",
     "2020 — 2025"),
]
for t, d in teach:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(t); _set_run(r, 9)
    r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

# ---------------------------------------------------------------- Skills
section("Technical Skills & Certifications")
skills = [
    ("Programming & analysis", "Python, R, Google Earth Engine"),
    ("LiDAR / point cloud", "ALS · GEDI · MLS/TLS processing, LAStools / PDAL, canopy & vertical-structure metrics"),
    ("Remote sensing & GIS", "QGIS, ArcGIS, satellite & UAV imagery, thermal imaging, Envi-met microclimate"),
    ("Machine learning", "classification / regression (XGBoost, deep learning) for forest & wildlife mapping"),
    ("Certifications", "Engineer Forest, Engineer Landscape Architecture, Drone Pilot License (Class 2)"),
    ("Languages", "Korean (native), English (professional working), Chinese (elementary)"),
]
for label, val in skills:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(label + ":  "); _set_run(r, 9, bold=True, color=ACCENT)
    r = p.add_run(val); _set_run(r, 9)

# ---------------------------------------------------------------- References
section("References")
refs = [
    ("Prof. Youngkeun Song (Doctoral advisor)",
     "Landscape & Ecological Planning Lab, Seoul National University · [email — to be added]"),
    ("Prof. Brady S. Hardiman",
     "Dept. of Forestry & Natural Resources, Purdue University · [email — to be added]"),
    ("Prof. Jaeyoung Ha",
     "School of Architecture + Design, Virginia Tech · [email — to be added]"),
]
for name, detail in refs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(name + " — "); _set_run(r, 9, bold=True)
    r = p.add_run(detail); _set_run(r, 9, color=GREY)

# ---------------------------------------------------------------- save
out = "Lee_Seunghyeon_CV_v1.docx"
doc.save(out)
print("saved:", out)
