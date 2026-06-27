# -*- coding: utf-8 -*-
"""
Academic CV builder — Seunghyeon (Clay) Lee
Forest remote sensing / LiDAR · GEDI / ecology — postdoc application

Structure follows field-standard academic CVs (Pascual UMD / Harvard GSAS).
All content is sourced from the verified portfolio data in
  ../src/data/site.json
Edit the data blocks below, then run:  python build_cv.py
Output: Lee_Seunghyeon_CV_v1.docx  ->  export to PDF in Word for ../public/cv/
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, json
from datetime import date as _dtdate

HERE = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------- palette
INK    = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x1f, 0x4e, 0x3d)   # deep forest green
GREY   = RGBColor(0x55, 0x55, 0x55)
LIGHT  = RGBColor(0x88, 0x88, 0x88)
BODYFONT = "Cambria"   # academic serif, high legibility at small sizes; clean via LibreOffice
HEADFONT = "Cambria"

doc = Document()

# page margins
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

RIGHT_TAB = Inches(6.9)  # within 7.0" text width

# Disable OpenType kerning/ligatures — Word applies them unevenly when exporting to PDF,
# producing the "some letters too tight, some too loose" spacing. Turning the feature off
# gives uniform metric spacing.
for _cs in doc.settings.element.iter(qn('w:compatSetting')):
    if _cs.get(qn('w:name')) == 'enableOpenTypeFeatures':
        _cs.set(qn('w:val'), '0')

# base style
st = doc.styles["Normal"]
st.font.name = BODYFONT
st.font.size = Pt(9.5)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.04

def _force_font(rPr, font):
    """Set the font for ALL script categories (ascii/hAnsi/cs/eastAsia) so Word never
    falls back to the theme font for punctuation (·, —, “”, †, →) — that fallback splits
    the text across font subsets and produces uneven letter/word spacing in the PDF."""
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for _a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(_a), font)

def _set_run(r, size=9.5, bold=False, italic=False, color=INK, font=BODYFONT, caps=False, spacing=None):
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    rPr = r._element.get_or_add_rPr()
    _force_font(rPr, font)
    if spacing is not None:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing)); rPr.append(sp)

def add_hyperlink(paragraph, url, text, color="555555", size=9,
                  underline=False, italic=False, bold=False, font=BODYFONT):
    """Append a real, clickable hyperlink run to a paragraph (works in Word and exported PDF)."""
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), font)
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if italic: rPr.append(OxmlElement('w:i'))
    if bold:   rPr.append(OxmlElement('w:b'))
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t); link.append(run); paragraph._p.append(link)
    return link

def add_right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

def bottom_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1f4e3d')
    pbdr.append(b); pPr.append(pbdr)

# ---------------------------------------------------------------- running header / footer (every page)
_sec = doc.sections[0]
_hp = _sec.header.paragraphs[0]
_hp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
_set_run(_hp.add_run("\tSeunghyeon Lee  —  Curriculum Vitae"), 8, color=LIGHT)
_fp = _sec.footer.paragraphs[0]
_fp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
_set_run(_fp.add_run("\tSeunghyeon Lee    "), 8, color=LIGHT)
_pg = _fp.add_run()
for _t, _v in (("begin", None), ("instr", " PAGE "), ("end", None)):
    if _t == "instr":
        e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = _v
    else:
        e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), _t)
    _pg._r.append(e)
_set_run(_pg, 8, color=LIGHT)

# ---------------------------------------------------------------- header
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Seunghyeon Lee"); _set_run(r, 22, bold=True, color=INK)
r = p.add_run("  (Clay Lee)"); _set_run(r, 12, color=GREY)
r = p.add_run("    ·    Curriculum Vitae"); _set_run(r, 10.5, color=LIGHT, caps=True, spacing=20)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Ph.D., expected Aug 2026 · Forest Structure & Remote Sensing & GIS & AI")
_set_run(r, 10, color=ACCENT, bold=True)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Landscape & Ecological Planning Lab, Seoul National University, Seoul, Republic of Korea")
_set_run(r, 9.5, color=GREY)
p.paragraph_format.space_after = Pt(2)

contact = [
    ("shlee5598@snu.ac.kr", "mailto:shlee5598@snu.ac.kr"),
    ("+82-10-8455-5598", None),
    ("ORCID 0000-0003-4612-065X", "https://orcid.org/0000-0003-4612-065X"),
    ("Google Scholar", "https://scholar.google.com/citations?user=Ew1_-r0AAAAJ"),
    ("drseunghyeonlee.com", "https://drseunghyeonlee.com/en"),
]
p = doc.add_paragraph()
for i, (disp, url) in enumerate(contact):
    if i:
        sep = p.add_run("  |  "); _set_run(sep, 9, color=LIGHT)
    if url:
        add_hyperlink(p, url, disp, color="555555", size=9)
    else:
        _set_run(p.add_run(disp), 9, color=GREY)
bottom_border(p); p.paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------------- section helper
def section(title):
    p = doc.add_paragraph()
    r = p.add_run(title); _set_run(r, 11, bold=True, color=ACCENT, caps=True, spacing=30)
    bottom_border(p)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)

def entry(left_bold, right="", sub="", bullets=None, left_rest=""):
    """One CV entry: bold lead + optional right-aligned date, italic subline, bullets."""
    p = doc.add_paragraph(); add_right_tab(p)
    r = p.add_run(left_bold); _set_run(r, 9.5, bold=True)
    if left_rest:
        r = p.add_run(left_rest); _set_run(r, 9.5)
    if right:
        r = p.add_run("\t" + right); _set_run(r, 9, color=GREY)
    p.paragraph_format.space_before = Pt(3)
    if sub:
        ps = doc.add_paragraph(); r = ps.add_run(sub); _set_run(r, 9, italic=True, color=GREY)
    if bullets:
        for b in bullets:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent = Inches(0.25)
            pb.paragraph_format.space_after = Pt(0)
            r = pb.add_run(b); _set_run(r, 9)

def authors_run(p, authors, me="Seunghyeon Lee", mark_first=False):
    """Render byline; bold `me`. If mark_first and `me` is the first author, add † after the name."""
    for i, a in enumerate(authors):
        if i:
            s = p.add_run(", "); _set_run(s, 9, color=INK)
        r = p.add_run(a); _set_run(r, 9, bold=(a == me))
        if mark_first and i == 0 and a == me:
            d = p.add_run("†"); _set_run(d, 9, bold=True, color=ACCENT)

# ---------------------------------------------------------------- Research Interests
section("Research Interests")
p = doc.add_paragraph()
r = p.add_run("Airborne & spaceborne LiDAR (GEDI) for vertical forest structure · forest typology and "
              "stratification · biomass and carbon estimation · urban ecosystems and green infrastructure · "
              "reproducible, code-driven remote-sensing science.")
_set_run(r, 9.5)
p.paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- Education
section("Education")
edu = [
    ("Ph.D., Interdisciplinary Program in Landscape Architecture",
     "2022.03 — Aug 2026",
     "Integrated Major in SmartCity Global Convergence · Seoul National University · "
     "Landscape & Ecological Planning Lab · Advisor: Prof. Youngkeun Song",
     ["Dissertation: “LiDAR-Based Characterization of Vertical Structure and Disturbance in "
      "Temperate Forests”",
      "Committee: Dongkun Lee, Youngryel Ryu (Seoul National Univ.); Hansoo Kim (Gyeonggi Research "
      "Institute); Dennis Heejoon Choi (Dankook Univ.); Youngkeun Song (advisor, SNU)"]),
    ("M.A., Landscape Architecture",
     "2020.03 — 2022.02",
     "Integrated Major in SmartCity Global Convergence · Graduate School of Environmental Studies, "
     "Seoul National University · Advisor: Prof. Youngkeun Song",
     ["Thesis: “Feasibility of Wildlife Detection Using UAV-derived Thermal and True-color Imagery”"]),
    ("B.S., Landscape Architecture",
     "2011.03 — 2017.02",
     "College of Agriculture & Life Sciences, Kyungpook National University, Daegu, Republic of Korea", None),
]
for lead, date, sub, bl in edu:
    entry(lead, right=date, sub=sub, bullets=bl)

# ---------------------------------------------------------------- Research Experience
section("Research Experience")
exp = [
    ("Visiting Research Intern", " — Purdue University, FNR FUSE Lab (PI: Prof. Brady Hardiman)",
     "2023.12 — 2024.02", "West Lafayette, IN, United States · Forest structure & LiDAR"),
    ("Visiting Research Intern", " — Virginia Tech (PI: Prof. Jaeyoung Ha)",
     "2024.07 — 2024.08", "Blacksburg, VA, United States · Landscape architecture"),
    ("Graduate Researcher", " — Landscape & Ecological Planning Lab, Seoul National University",
     "2020.03 — Present", "National R&D and municipal projects on LiDAR forest structure, carbon, and urban ecology"),
    ("Research Assistant", " — Architecture & Urban Research Institute (AURI)",
     "2019.08 — 2019.11", "Republic of Korea"),
]
for lead, rest, date, sub in exp:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- Publications
section("Peer-Reviewed Publications")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  "
              "Impact Factor / quartile shown where available.  (SCI then KCI; newest first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

# --- source: citations/publications.bib (ascending) + IF/quartile from site.json (matched by DOI) ---
def _bib_entries(path):
    txt = open(path, encoding="utf-8").read()
    out, i = [], 0
    while True:
        at = txt.find("@", i)
        if at < 0: break
        b = txt.find("{", at); depth, j = 0, b
        while j < len(txt):
            if txt[j] == "{": depth += 1
            elif txt[j] == "}":
                depth -= 1
                if depth == 0: break
            j += 1
        out.append(txt[at:j + 1]); i = j + 1
    return out

def _bf(raw, name):
    m = re.search(r"\b" + name + r"\s*=\s*\{(.*?)\}\s*[,}\n]", raw, re.I | re.S)
    return m.group(1).strip() if m else ""

def _initials(given):
    return " ".join("-".join(s[0].upper() + "." for s in part.split("-") if s)
                    for part in given.split())

def _apa_author(name):
    if "," in name:
        sur, giv = [x.strip() for x in name.split(",", 1)]
    else:
        ps = name.split(); sur, giv = ps[-1], " ".join(ps[:-1])
    return f"{sur}, {_initials(giv)}".strip().rstrip(","), (sur == "Lee" and giv.startswith("Seunghyeon"))

_site = json.load(open(os.path.join(HERE, "..", "src", "data", "site.json"), encoding="utf-8"))
IFQ = {(it.get("doi") or "").lower().strip(): (it.get("if", ""), it.get("quartile", ""))
       for it in _site["publications"]["items"] if it.get("doi")}

_MO = {m: i for i, m in enumerate(
    ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"], 1)}
def _monthnum(raw):
    m = _bf(raw, "month").lower()[:3]
    return _MO.get(m, int(m) if m.isdigit() else 0)

_pubs = []
for raw in _bib_entries(os.path.join(HERE, "citations", "publications.bib")):
    doi = _bf(raw, "doi")
    iff, q = IFQ.get(doi.lower(), ("", ""))
    yv = _bf(raw, "year"); yv = int(yv) if yv.isdigit() else 0
    _pg = re.search(r"\d+", _bf(raw, "pages")); sp = int(_pg.group()) if _pg else 0
    _pubs.append({
        "apa": [_apa_author(a.strip()) for a in _bf(raw, "author").split(" and ")],
        "year": _bf(raw, "year"),
        "title": re.sub(r"\s*-\s*$", "", _bf(raw, "title")).replace(" - ", ": ").rstrip(". "),
        "journal": _bf(raw, "journal").replace("\\&", "&"),
        "vol": _bf(raw, "volume"), "num": _bf(raw, "number"),
        "pages": _bf(raw, "pages").replace("-", "–"), "doi": doi, "iff": iff, "q": q,
        "sortk": (0 if (iff or q) else 1, -yv, -_monthnum(raw), sp),  # SCI→KCI, newest first, then start page
    })
_pubs.sort(key=lambda d: d["sortk"])

for n, d in enumerate(_pubs, 1):
    apa, year, title, journal = d["apa"], d["year"], d["title"], d["journal"]
    vol, num, pages, doi, iff, q = d["vol"], d["num"], d["pages"], d["doi"], d["iff"], d["q"]
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({year}). "), 9)
    _set_run(p.add_run(title + ". "), 9)
    _set_run(p.add_run(journal), 9, italic=True)
    if vol:
        _set_run(p.add_run(", "), 9)
        _set_run(p.add_run(vol), 9, italic=True)
        if num:
            _set_run(p.add_run(f"({num})"), 9)
        if pages:
            _set_run(p.add_run(f", {pages}"), 9)
    _set_run(p.add_run(". "), 9)
    if doi:
        add_hyperlink(p, f"https://doi.org/{doi}", f"https://doi.org/{doi}", color="555555", size=9)
    if iff or q:
        iv = iff.split("(")[0].strip()
        tag = "  · " + ", ".join(x for x in [f"IF {iv}" if iv else "", q] if x)
        _set_run(p.add_run(tag), 8.5, color=ACCENT)

# ---------------------------------------------------------------- Manuscripts under review
section("Manuscripts Under Review")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  (5 manuscripts; first-author listed first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
# status = "stage|date"  (stage shown in brackets, date in grey)
under_review = [
    (["Seunghyeon Lee", "Dennis Heejoon Choi", "Youngkeun Song", "James H. Thorne"],
     "Spaceborne LiDAR reveals post-fire vertical structural loss in a dense temperate Asian conifer "
     "plantation.", "Science of Remote Sensing", "Manuscript under review|2026.04"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"],
     "Systematic bias in urban-forest vegetation coverage assessment: the influence of topo-edaphic "
     "conditions on discrepancies between ALS and field surveys.", "Geoscience Letters",
     "Revise and resubmit, 1st round|2026.06"),
    (["Seunghyeon Lee", "Yonghwan Kim", "Dohee Kim", "Hansoo Kim", "Youngkeun Song"],
     "Bi-temporal ALS assessment of vertical–horizontal forest structures and their "
     "association in a temperate urban forest.", "Forest Science and Technology",
     "Revise and resubmit, 2nd round|2026.05"),
    (["Yonghwan Kim", "Seunghyeon Lee", "Wonhyeop Shin", "Youngkeun Song"],
     "Integrating UAV-derived habitat metrics with movement persistence and species distribution models "
     "to characterize seasonal habitat use of invasive turtles in urban wetlands.",
     "Global Ecology and Conservation", "Revise and resubmit, 1st round|2026.05"),
    (["Gapseong Jekal", "Yong Hwan Kim", "Seunghyeon Lee", "Ji Weon Yun", "Dae Yeol Kim", "Youngkeun Song"],
     "Monitoring transition-zone dynamics of a Phragmites communis–Suaeda japonica mosaic from "
     "multi-season Sentinel-2 in a coastal wetland.", "Journal of Coastal Conservation",
     "Revise and resubmit, 1st round|2026.02"),
]
for n, (auth, title, jour, status) in enumerate(under_review, 1):
    apa = [_apa_author(a) for a in auth]
    stage, _, dt = status.partition("|")
    yr = dt[:4] if dt[:4].isdigit() else "2026"
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({yr}). "), 9)
    _set_run(p.add_run(title.rstrip(". ")), 9, italic=True)       # APA: unpublished title in italics
    _set_run(p.add_run(f" [{stage}]. "), 9)
    _set_run(p.add_run(jour + "."), 9)
    if dt:
        _set_run(p.add_run(f"  · {dt}"), 8.5, color=GREY)

# ---------------------------------------------------------------- Books
section("Books")
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
_set_run(p.add_run("1. "), 9, color=GREY)
_set_run(p.add_run("YOZMDOSI collective ("), 9)
_set_run(p.add_run("Seunghyeon Lee"), 9, bold=True)
_set_run(p.add_run(", contributing author) (2021). "), 9)
_set_run(p.add_run("New Normal City"), 9, italic=True)
_set_run(p.add_run(". Urban Issue. ISBN 9791197343100."), 9)

# ---------------------------------------------------------------- Fellowships & Funding
section("Fellowships & Funding")
p = doc.add_paragraph()
r = p.add_run("Competitive fellowships / scholarships (amounts converted at ₩1,400 = $1).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
funding = [
    ("Chung Mong-Koo Foundation — OnDream Future Industry Talent Graduate Fellowship (full)",
     "2022 — 2025", "$25,000"),
    ("BK21 FOUR — Integrated Major in SmartCity Global Convergence (full)", "2020 — 2025", "$39,285"),
    ("Ilju Foundation Undergraduate Scholarship, 23rd Scholar (full)", "2015 — 2017", "$10,714"),
    ("Challenge Scholarship, Kyungpook National University (full)", "2011 — 2013", "$18,571"),
]
for name, period, amount in funding:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
    _set_run(p.add_run(name), 9)
    _set_run(p.add_run(f"  ({amount})"), 9, color=ACCENT)
    _set_run(p.add_run("\t" + period), 9, color=GREY)

# ---------------------------------------------------------------- Invited talks
section("Invited Talks")
invited = [
    ("The understanding and applications of thermal sensors.",
     "Guest lecture, SmartCity Industrial Technology Seminar (3 cr.), Seoul National University", "2025"),
    ("Understanding spatial information and research cases.",
     "Guest lecture, Environmental Conservation & Land Management (3 cr.), Seoul National University", "2025"),
    ("Remotely sensed smart-city ecological value maps.",
     "NEF (Neocity Empowerment Forum), Orlando, FL", "2024"),
    ("Invasive species monitoring development and its application.",
     "Asia Week, Fukuoka, Japan", "2024"),
]
for title, venue, yr in invited:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    b = p.add_run("• "); _set_run(b, 9, color=ACCENT)
    r = p.add_run("Seunghyeon Lee"); _set_run(r, 9, bold=True)
    r = p.add_run(f" ({yr}). "); _set_run(r, 9)
    r = p.add_run(title + " "); _set_run(r, 9)
    r = p.add_run(venue + "."); _set_run(r, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Conferences
section("Conference Presentations (International)")
p = doc.add_paragraph()
r = p.add_run("First author / presenter unless marked “team”.  "
              "10 domestic presentations (KOSERT, KSEE, KILA) available on request.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

confs = [
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2025, "Poster",
     "Novel LiDAR indices reveal scale-dependent vertical structure and typologies in a temperate forest.",
     "AGU Fall Meeting, New Orleans, US"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Poster",
     "Quantifying the number of forest stratification layers of city-scale forest and characterizing by "
     "forest types.", "AGU Fall Meeting, Washington D.C., US"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Oral",
     "Canopy layers distribution by forest patch and species at the city-level forest using airborne LiDAR.",
     "ICLEE, Kitakyushu, Japan"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Poster",
     "Characterizing forest vegetation stratification by forest biotope types using airborne LiDAR.",
     "ESA Annual Meeting, Long Beach, CA, US"),
    (["BK21 SmartCity Team"], 2024, "Video, team",
     "Smart city & green infrastructure.", "CES (Consumer Electronics Show), Las Vegas, NV"),
    (["Landscape & Ecological Planning Lab"], 2023, "Oral, team",
     "Property-information analysis of Gyeonggi-do Province biotopes.",
     "JCK Symposium (18th), Kyoto, Japan"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2023, "Poster",
     "Wildfire-driven forest vegetation height change comparison.", "JCK Symposium (18th), Kyoto, Japan"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2023, "Poster",
     "The impacts of fire-induced disturbances on tree height and structure using GEDI-derived variables.",
     "ESA Annual Meeting, Portland, OR, US"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Monitoring structural change of fire-induced forest vegetation using GEDI.", "ICLEE (12th), Online"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Comparison of GEDI and aerial laser scanning datasets according to leaf-on and leaf-off seasons.",
     "AGU Fall Meeting, Chicago, IL, US / Online"),
    (["Seunghyeon Lee", "Dennis Heejoon Choi", "Youngkeun Song"], 2022, "Video",
     "Classification of tree species and forest floor using XGBoost with forest maps, "
     "airborne LiDAR, and satellite imagery.", "World Forestry Congress (15th), Seoul, South Korea"),
    (["Seunghyeon Lee", "Sung-Ho Kil", "Youngkeun Song"], 2022, "Poster",
     "Feasibility analyses of real-time detection of wildlife using UAV-derived thermal and RGB images.",
     "World Forestry Congress (15th), Seoul, South Korea"),
    (["Seunghyeon Lee", "DaeYeol Kim", "Dennis Heejoon Choi", "Hansoo Kim", "Youngkeun Song"], 2021, "Poster",
     "Detecting individual broad-leaved trees by a trunk-extraction method using leaf-off airborne LiDAR.",
     "AGU Fall Meeting, New Orleans, US / Online"),
]
for auth, yr, kind, title, venue in confs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
    authors_run(p, auth)
    yr_r = p.add_run(f" ({yr}). "); _set_run(yr_r, 9)
    t = p.add_run(title + " "); _set_run(t, 9)
    v = p.add_run(f"[{kind}] {venue}."); _set_run(v, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Grants / Projects
section("Research Projects")
p = doc.add_paragraph()
r = p.add_run("National R&D and municipal projects (newest first).  "
              "Budget = total project funding (national R&D).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

projects = [
    ("Janghang National Wetland Restoration — ecosystem monitoring",
     "Korea Environmental Conservation Institute", "2025.07 — 2025.11", "Researcher",
     "Prof. Youngkeun Song (SNU)", None,
     ["Led UAV multispectral and LiDAR mapping of the restoration site",
      "Estimated current carbon stock from UAV LiDAR and predicted future stock from the wetland-park vegetation master plan"],
     "Coastal-wetland multi-season Sentinel-2 monitoring (manuscript under review)"),
    ("Carbon storage/sink analysis and inventory for Gyeonggi-do by biotope type",
     "Gyeonggi Research Institute", "2024.06 — 2025.05", "Researcher",
     "Prof. Youngkeun Song (SNU)", None,
     ["Analyzed province-wide vertical and horizontal forest structure from ALS",
      "Estimated forest carbon storage by fusing ALS, airborne hyperspectral, and PlanetScope imagery"],
     "Multi-scale forest typologies (Ecological Indicators 2026, first author); post-fire structural loss via spaceborne LiDAR (Science of Remote Sensing, under review); coastal-wetland Phragmites–Suaeda monitoring (Journal of Coastal Conservation, under review); AGU & ESA presentations"),
    ("Integrated assessment tool for carbon reduction by offset and synergy with ecosystem services",
     "Korean Ministry of Environment", "2023.04 — 2027.12", "Researcher",
     "Prof. Youngkeun Song (SNU)", "$1.13M",
     ["Acquired 3-D structure of abandoned-paddy wetlands using drone LiDAR",
      "Quantified carbon storage of abandoned paddies and adjacent vegetation"],
     "Insect diversity & ecological structure (Agric., Ecosyst. & Environ. 2026); Rook habitat preferences (Global Ecology & Conservation 2025)"),
    ("Gyeonggi-do biotope property-information analysis",
     "Gyeonggi Research Institute", "2022.12 — 2023.11", "Researcher",
     "Prof. Youngkeun Song (SNU)", None,
     ["Developed and quantified a method to analyze vegetation layer structure by forest type"],
     "Forest layer-typology metrics (Ecological Indicators 2026); JCK 2023 presentation"),
    ("Creation, restoration & management technology of carbon-accumulated abandoned-paddy wetland",
     "Korean Ministry of Environment", "2022.04 — 2026.12", "Researcher",
     "Prof. Bonhak Koo (Sangmyung Univ.)", "$391K",
     ["Mapped abandoned-paddy sites and built field drone multispectral / LiDAR datasets"],
     "Non-destructive carbon storage of Salix spp. (KOSERT 2025); paddy-terrestrialization patent (pending)"),
    ("Gwacheon-si urban ecological status mapping",
     "Gyeonggi Research Institute", "2021.09 — 2022.12", "Researcher",
     "Prof. Youngkeun Song (SNU)", None,
     ["Classified forest types and tree species of Gwacheon-si forests and urban green spaces",
      "Analyzed the vertical and horizontal forest structure of Gwacheon-si"],
     "Forest-type GEDI–ALS canopy height (KOSERT 2026, first author); systematic ALS-vs-field urban-forest bias (Geoscience Letters, under review); bi-temporal ALS forest-structure assessment (Forest Science & Technology, under review)"),
    ("Real-time web-based positioning surveillance system for introduced exotic species",
     "Korean Ministry of Environment", "2021.04 — 2023.12", "Project Manager",
     "Prof. Youngkeun Song (SNU)", "$863K",
     ["Led the project as managing researcher — scheduling, budget, and research-output management",
      "Attached GPS trackers to four invasive species (red-eared slider, Formosan sika deer, nutria, raccoon) and built movement datasets",
      "Built drone multispectral, LiDAR, and hyperspectral imagery of invasive-turtle habitats",
      "Produced invasive-species control manuals and developed a real-time positioning-tracking platform"],
     "Invasive-turtle habitat model (manuscript under review, GEC); wildlife-tracker missing-point patent (reg. 2023) and optimal-capture-range patent (pending)"),
    ("IT·ET·BT convergence-based distribution and spread models for introduced exotic species",
     "Korean Ministry of Environment", "2021.04 — 2023.12", "Researcher",
     "Prof. Wanmo Kang (Kookmin Univ.)", "$464K",
     ["Modeled habitat distribution and spread of invasive turtles from GPS-tracking data using MaxEnt"],
     "Invasive-turtle MaxEnt distribution model (manuscript under review, GEC)"),
    ("Redundancy-based green-infrastructure technology for urban ecosystem challenges",
     "Korean Ministry of Environment", "2020.04 — 2022.12", "Researcher",
     "Prof. Youngkeun Song (SNU)", "$662K",
     ["Tracked wild-boar distribution / movement with camera traps and analyzed home range",
      "Built drone multispectral / LiDAR imagery of wild-boar habitat and performed real-time thermal-UAV tracking"],
     "Diel activity of water deer & wild boar via long-term camera-trapping (KOSERT 2024); "
     "drone thermal/RGB & LiDAR wild-boar habitat datasets"),
    ("Monitoring and simulation of climate conditions by spatial type",
     "The Seoul Institute", "2020.01 — 2020.08", "Researcher",
     "Prof. Youngkeun Song (SNU)", None,
     ["Acquired and analyzed field measurements of particulate matter and temperature/humidity by urban spatial type",
      "Ran ENVI-met microclimate simulations and designed green-space planting plans for PM reduction, predicting post-implementation outcomes"],
     "LST estimation comparison — satellite / simulation / UAV (KOSERT 2026); outdoor thermal-comfort optimization patent (reg. 2022)"),
    ("Customized habitat-management technique for urban species",
     "Korean Ministry of Environment", "2019.04 — 2022.12", "Researcher",
     "Prof. Chan Park (Univ. of Seoul)", "$525K",
     ["Developed a wildlife-detection method using thermal UAV imagery"],
     "Real-time UAV thermal/RGB wildlife detection (Remote Sensing 2021, first author); UAV thermal wildlife-detection patent (reg. 2022)"),
    ("Assessment scheme and ecosystem-restoration model by degradation type",
     "Korean Ministry of Environment", "2018.07 — 2020.12", "Researcher",
     "Prof. Youngkeun Song (SNU)", "$257K",
     ["Performed land-cover classification using satellite imagery"],
     None),
]
for title, funder, period, role, pi, budget, bullets, outputs in projects:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
    _set_run(p.add_run(title), 9.5, bold=True)
    _set_run(p.add_run("\t" + period), 9, color=GREY)
    ps = doc.add_paragraph()
    _set_run(ps.add_run(f"{funder} · "), 9, italic=True, color=GREY)
    _set_run(ps.add_run(f"Role: {role}"), 9, italic=True, bold=True, color=ACCENT)
    _set_run(ps.add_run(f" · PI: {pi}"), 9, italic=True, color=GREY)
    if budget:
        _set_run(ps.add_run(f" · {budget} (total project)"), 9, italic=True, color=GREY)
    for b in bullets:
        pb = doc.add_paragraph(style="List Bullet")
        pb.paragraph_format.left_indent = Inches(0.25); pb.paragraph_format.space_after = Pt(0)
        _set_run(pb.add_run(b), 9)
    if outputs:
        po = doc.add_paragraph(); po.paragraph_format.left_indent = Inches(0.25)
        po.paragraph_format.space_before = Pt(1)
        _set_run(po.add_run("→ Related outputs:  "), 8.5, bold=True, color=ACCENT)
        _set_run(po.add_run(outputs), 8.5, italic=True, color=GREY)

# ---------------------------------------------------------------- Patents
section("Patents (Republic of Korea)")
p = doc.add_paragraph()
r = p.add_run("7 registered, 4 pending.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

def _patent_sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _patent_list(items):
    for title, meta in items:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
        t = p.add_run(title + " "); _set_run(t, 9)
        m = p.add_run(meta); _set_run(m, 8.5, italic=True, color=GREY)

_patent_sub("Registered")
_patent_list([
    ("System and method for providing a tree-management platform.", "Registered 2025.09 · Sole inventor"),
    ("Method and system for determining wetland vegetation structure using drone LiDAR.",
     "Registered 2025.04 · Seoul National University"),
    ("Image-based roadside-tree information management system and method.", "Registered 2024.11 · Sole inventor"),
    ("Prediction of missing location points of a wildlife GPS tracker by combining in-situ and biological "
     "information.", "Registered 2023.07 · Seoul National University"),
    ("Self-heating target module for thermal-camera performance testing and precision data acquisition.",
     "Registered 2022.10 · Seoul National University"),
    ("Automated wildlife detection method and device using drone-mounted thermal-camera imagery.",
     "Registered 2022.05 · Seoul National University"),
    ("A method to derive an optimal plan for outdoor thermal-comfort mitigation.",
     "Registered 2022.04 · Seoul National University"),
])
_patent_sub("Pending")
_patent_list([
    ("Quantitative assessment and automatic classification of wildfire damage using spaceborne LiDAR.",
     "Pending 2025 · Seoul National University & Dankook Univ."),
    ("Monitoring halophyte distribution from multi-period vegetation indices in coastal wetlands.",
     "Pending 2025 · Seoul National University"),
    ("Determining terrestrialization of abandoned-paddy wetlands using multi-temporal drone LiDAR and "
     "stratigraphic volume change.", "Pending 2023 · Seoul National University"),
    ("Deriving optimal wildlife-capture range via machine learning by fusing GPS-tracking data and drone "
     "footage.", "Pending 2023 · Seoul National University"),
])

# ---------------------------------------------------------------- Honors
section("Awards & Honors")

def _subhead(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _tabbed(rows):
    for h, d in rows:
        p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h); _set_run(r, 9)
        r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

_subhead("Awards")
_tabbed([
    ("Best Presentation Award — KOSERT Autumn Meeting "
     "(GEDI vs. airborne LiDAR field agreement by vegetation type and season)", "2022"),
    ("Best Presentation Award — KOSERT Spring Meeting "
     "(forest layer-structure indices and their correlation with carbon storage)", "2025"),
])

# ---------------------------------------------------------------- Service & Membership
section("Professional Service & Membership")
p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
_set_run(p.add_run("Peer reviewer, "), 9, bold=True, color=ACCENT)
_set_run(p.add_run("International Journal of Applied Earth Observation and Geoinformation"), 9)
_set_run(p.add_run("\t2026 – present"), 9, color=GREY)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
_set_run(p.add_run("Professional memberships"), 9, bold=True, color=ACCENT)
members = [
    ("American Geophysical Union (AGU)", "2021 – present"),
    ("Ecological Society of America (ESA)", "2021 – present"),
    ("Korean Society of Environmental Restoration Technology (KOSERT)", "2020 – present"),
    ("Ecological Society of Korea (ESK)", "2023 – present"),
]
for name, period in members:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(1)
    _set_run(p.add_run("Member, " + name), 9)
    _set_run(p.add_run("\t" + period), 9, color=GREY)

# ---------------------------------------------------------------- Teaching
section("Teaching Experience")
p = doc.add_paragraph()
r = p.add_run("2 university courses · 4 online courses (1,000+ students) · 20+ invited lectures and workshops.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

def _t_sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _t_item(text, date, url=None):
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    b = p.add_run("• "); _set_run(b, 9, color=ACCENT)
    if url:
        add_hyperlink(p, url, text, color="1a1a1a", size=9)   # clickable title, normal colour
    else:
        r = p.add_run(text); _set_run(r, 9)
    r = p.add_run("\t" + date); _set_run(r, 9, color=GREY)

_t_sub("University Courses")
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
_set_run(p.add_run("• "), 9, color=ACCENT)
_set_run(p.add_run("Instructor — “Understanding and Application of Spatial Information” "
                   "(3 credits; 30 students)"), 9)
p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.28)
_set_run(p.add_run("Division of Architecture & Urban Design, Incheon National University."), 9, color=GREY)
p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.28)
_set_run(p.add_run("Designed and taught a semester-long undergraduate course (20% lecture / 80% hands-on lab): "
    "GIS fundamentals (coordinate systems, vector/raster data models); QGIS labs (data acquisition, editing, "
    "spatial analysis, cartographic visualization); and remote sensing — satellite-image preprocessing and "
    "spectral-index (e.g., NDVI) analysis, plus UAV imagery."), 9, color=GREY)
p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.left_indent = Inches(0.28)
_set_run(p.add_run("Term project: public open-data GIS analysis projects and map outputs on urban, "
                   "architectural, environmental, and ecological topics for the Incheon region."), 9, color=GREY)
_set_run(p.add_run("\t2024 — 2025"), 9, color=GREY)

_t_sub("Online Courses  (1,000+ students enrolled)")
_t_item("QGIS Trendy Visualization — Election-Result Mapping", "2025", "https://inf.run/Vo32x")
_t_item("QGIS Beginner All-in-One Starter Pack (theory & practice)", "2024", "https://inf.run/cunhy")
_t_item("QGIS Mapping Visualization A to Z (Vector / Basic)", "2023", "https://inf.run/JcRA")
_t_item("QGIS Python Automation", "2022", "https://inf.run/RNcr")

_t_sub("Invited Lectures & Workshops  (20+ sessions, 2020 — 2025)")
_t_item("Hyundai NGV — Understanding GIS & spatial analysis/visualization via DBMS (basic & advanced)", "2025")
_t_item("Korea Water Resources Corp. (K-water) — Q-GIS practice for water-sector professionals", "2023 — 2025")
_t_item("Korea Environment Corporation (K-eco) — Data-analysis expert training program", "2024 — 2025")
_t_item("Seoul Software Academy (SeSAC) — AI Big-Data Analyst course: GIS module", "2025")
_t_item("Chungcheongnam-do Agricultural Research & Extension Services — QGIS disease-spread visualization", "2024")
_t_item("Gyeonggi-do HRD Institute — Data visualization for better reports", "2022")
_t_item("BK21 SmartCity · SNU GSES — WISE-UP urban spatial-data analysis workshop", "2024")
_t_item("SNU Lifelong Education Center — “Map My School Using Public Data” (high-school outreach)", "2024")

# ---------------------------------------------------------------- Professional Experience
section("Professional Experience")
prof = [
    ("Founder / CEO", " — TREE:ID (startup)", "2023.09 — 2025.03",
     "A web platform for individual-tree detection & mapping — urban-tree inventory, diagnosis, and "
     "management. ~$71,000 startup grant; two patents held as assignee. Republic of Korea"),
    ("Landscape Architect", " — SUNJIN Engineering & Architecture Co., Ltd. (full-time)",
     "2017.03 — 2018.11", "Seoul, Republic of Korea · landscape and urban-planning practice"),
    ("Landscape Design Intern", " — ATLAS Landscape Architecture, China", "2016.09 — 2017.02",
     "China · landscape design practice"),
]
for lead, rest, date, sub in prof:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- References
section("References")
references = [
    ("Prof. Youngkeun Song",
     "Doctoral advisor (Ph.D. & M.A.)",
     "Landscape & Ecological Planning Lab, Department of Landscape Architecture,\n"
     "Graduate School of Environmental Studies, Seoul National University",
     "songkoon@gmail.com"),
]
for name, role, affil, email in references:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
    _set_run(p.add_run(name), 9.5, bold=True)
    _set_run(p.add_run("\t" + role), 9, italic=True, color=ACCENT)
    for line in affil.split("\n"):
        pa = doc.add_paragraph(); pa.paragraph_format.left_indent = Inches(0.0)
        pa.paragraph_format.space_after = Pt(0)
        _set_run(pa.add_run(line), 9, color=GREY)
    pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(0)
    _set_run(pe.add_run("Email: "), 9, color=GREY)
    add_hyperlink(pe, "mailto:" + email, email, color="555555", size=9)

# (Technical Skills section intentionally omitted)

# ---------------------------------------------------------------- date / signature
# Drop a signature image at cv/assets/signature.png (transparent PNG, tightly cropped) to embed it;
# otherwise a blank signature line is printed.
_sig_img = os.path.join(HERE, "assets", "signature.png")
_today = _dtdate.today().strftime("%d %B %Y")
if os.path.exists(_sig_img):
    pim = doc.add_paragraph(); pim.paragraph_format.space_before = Pt(16)
    pim.add_run().add_picture(_sig_img, height=Inches(0.6))
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(1)
    _set_run(p.add_run("Seunghyeon Lee"), 9.5, bold=True)
    _set_run(p.add_run("\tDate: " + _today), 9, color=GREY)
else:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(18)
    _set_run(p.add_run("Seunghyeon Lee"), 9.5, bold=True)
    _set_run(p.add_run("\tSignature: ____________________     Date: " + _today), 9, color=GREY)

# ---------------------------------------------------------------- save
out = "Lee_Seunghyeon_CV_v7.docx"
doc.save(out)
print("saved:", out)
