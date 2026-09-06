# -*- coding: utf-8 -*-
"""
Academic CV builder — Seunghyeon (Clay) Lee
Forest remote sensing / LiDAR · GEDI / ecology — postdoc application

Structure follows field-standard academic CVs (Pascual UMD / Harvard GSAS).
All content is sourced from the verified portfolio data in
  ../src/data/site.json
Edit the data blocks below, then run:  python build_cv.py
Output: Lee_Seunghyeon_CV_v1.docx  ->  export to PDF in Word for ../public/cv/
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, json
from datetime import date as _dtdate

HERE = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------- data source
# 사실 데이터의 정본은 application_info.xlsx 다. 이 파일에 값을 직접 적지 않는다.
#   - 원자 필드로 CV 표기가 재현되면 규칙으로 조합한다
#   - 재현되지 않는 curated 표기는 xlsx 의 'CV 표기(EN)' / 'CV 서술(EN)' 컬럼을 그대로 쓴다
#   - 수록 대상은 xlsx 의 'CV 표시' / 'CV 구분' / 'CV 분류' 플래그로 고른다
import openpyxl as _oxl

_WB = _oxl.load_workbook(os.path.join(HERE, "application_info.xlsx"), data_only=True)


def sheet(name, where=None, order=None):
    """xlsx 시트를 dict 리스트로. where 로 거르고 order 로 정렬한다.
    엑셀의 행 순서에 의존하지 않도록 정렬 키를 명시하는 것을 원칙으로 한다."""
    ws = _WB[name]
    hdr = [str(c.value) for c in ws[1]]
    out = [dict(zip(hdr, r)) for r in ws.iter_rows(min_row=2, values_only=True)]
    out = [r for r in out if any(v is not None for v in r.values())]
    if where:
        out = [r for r in out if where(r)]
    if order:
        out.sort(key=order)
    return out


def years_only(period):
    """'2022/09 — 2025' -> '2022 — 2025' (CV 는 장학 기간을 연도로만 표기)"""
    return re.sub(r"(\d{4})/\d{2}", r"\1", str(period or "")).strip()


def _dot(s):
    """제목 끝의 마침표는 표기 규칙이므로 렌더러가 붙인다 (xlsx 에는 넣지 않는다)."""
    return str(s or "").rstrip(". ") + "."


def _ym(v):
    """'2024/07/22' -> '2024.07'  (CV 는 경력 기간을 연.월로 표기)"""
    m = re.match(r"(\d{4})/(\d{2})", str(v or ""))
    return f"{m.group(1)}.{m.group(2)}" if m else ""


def _work(kind):
    """Work & Internships 에서 CV 구분으로 골라 (직위, 소속, 기간, 서술) 튜플로."""
    out = []
    for r in sheet("Work & Internships", where=lambda r: r["CV 구분"] == kind,
                   order=lambda r: _desc(r["Start Date"])):
        end = _ym(r["End Date"]) or "Present"
        out.append((r["CV 직위(EN)"], r["CV 소속(EN)"],
                    f'{_ym(r["Start Date"])} — {end}', r["CV 서술(EN)"]))
    return out


# 문장형 표기에서 소문자로 내리면 안 되는 것: 약어 · 고유명사 · 상표
_KEEP_WORDS = ["LiDAR", "GEDI", "XGBoost", "UAV", "ALS", "MLS", "RGB", "DEM", "DTM", "CHM",
               "Sentinel-2", "Landsat", "PlanetScope", "Gyeonggi-do", "Korea", "Korean",
               "Japan", "China", "Seoul", "AGU", "ESA", "ICLEE", "JCK", "CES", "KOSERT"]
_KEEP = {re.sub(r"[^a-z0-9]", "", w.lower()): w for w in _KEEP_WORDS}

# 여러 단어로 된 고유명사 — 단어 단위 보호로는 잡히지 않아 변환 후 되돌린다
_KEEP_PHRASES = ["Gyeonggi-do Province", "Republic of Korea", "World Forestry Congress",
                 "Consumer Electronics Show", "Seoul National University"]


def sentence_case(title):
    """제목형 -> 문장형 (APA 7th). 첫 단어와 보호 목록만 대문자를 유지한다.
    토큰 전체를 먼저 보호 목록과 대조하므로 Gyeonggi-do / Sentinel-2 처럼
    하이픈이 든 고유명사가 쪼개지지 않는다."""
    def one(tok, first):
        m = re.match(r"^(\W*)(.*?)(\W*)$", tok, re.S)
        pre, mid, post = m.group(1), m.group(2), m.group(3)
        if not mid:
            return tok
        bare = re.sub(r"[^a-z0-9]", "", mid.lower())
        if bare in _KEEP:
            return pre + _KEEP[bare] + post
        out = []
        for i, part in enumerate(mid.split("-")):
            pb = re.sub(r"[^a-z0-9]", "", part.lower())
            if pb in _KEEP:
                out.append(_KEEP[pb])
            elif first and i == 0:
                out.append(part[:1].upper() + part[1:].lower())
            else:
                out.append(part.lower())
        return pre + "-".join(out) + post

    out = " ".join(one(t, i == 0) for i, t in enumerate(str(title or "").split(" ")))
    for ph in _KEEP_PHRASES:
        out = re.sub(re.escape(ph), ph, out, flags=re.I)
    return out


def _lower_first(s):
    """괄호 안 설명은 소문자로 시작한다. 단 약어·고유명사(GEDI 등)는 그대로 둔다."""
    s = str(s or "")
    head = re.sub(r"[^a-z0-9]", "", s.split(" ")[0].lower())
    return s if head in _KEEP else s[:1].lower() + s[1:]


def _desc(v):
    """문자열 날짜를 내림차순 정렬 키로. CV 는 모든 섹션이 최신순이다."""
    return tuple(-ord(c) for c in str(v or ""))

# ---------------------------------------------------------------- palette
INK    = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x1f, 0x4e, 0x3d)   # deep forest green
GREY   = RGBColor(0x44, 0x44, 0x44)
LIGHT  = RGBColor(0x6f, 0x6f, 0x6f)
BODYFONT = "Cambria"   # academic serif, high legibility at small sizes; clean via LibreOffice
HEADFONT = "Cambria"

doc = Document()

# page margins
for s in doc.sections:
    s.top_margin = Inches(0.75); s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

RIGHT_TAB = Inches(7.1)  # within 7.1" text width

# Disable OpenType kerning/ligatures — Word applies them unevenly when exporting to PDF,
# producing the "some letters too tight, some too loose" spacing. Turning the feature off
# gives uniform metric spacing.
for _cs in doc.settings.element.iter(qn('w:compatSetting')):
    if _cs.get(qn('w:name')) == 'enableOpenTypeFeatures':
        _cs.set(qn('w:val'), '0')

# base style
st = doc.styles["Normal"]
st.font.name = BODYFONT
st.font.size = Pt(10)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.13

# ---- font-size scale (single source of truth for the v8 readability pass) ----
# Maps every legacy point size used in the call sites below to its new value, so the
# whole document rescales consistently without touching ~100 individual calls.
#   8/8.5 = notes/meta · 9/9.5/10 = body & titles · 10.5 = "CV" tag · 11 = section head · 22 = name
_SZ = {8: 8, 8.5: 9, 9: 10, 9.5: 10.5, 10: 10.5, 10.5: 11, 11: 13.5, 12: 12, 22: 24}

def _force_font(rPr, font):
    """Set the font for ALL script categories (ascii/hAnsi/cs/eastAsia) so Word never
    falls back to the theme font for punctuation (·, —, “”, †, →) — that fallback splits
    the text across font subsets and produces uneven letter/word spacing in the PDF."""
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for _a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(_a), font)

def _set_run(r, size=9.5, bold=False, italic=False, color=INK, font=BODYFONT, caps=False, spacing=None):
    r.font.size = Pt(_SZ.get(size, size)); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    rPr = r._element.get_or_add_rPr()
    _force_font(rPr, font)
    if spacing is not None:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing)); rPr.append(sp)

def add_hyperlink(paragraph, url, text, color="444444", size=9,
                  underline=False, italic=False, bold=False, font=BODYFONT):
    """Append a real, clickable hyperlink run to a paragraph (works in Word and exported PDF)."""
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), font)
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(_SZ.get(size, size) * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if italic: rPr.append(OxmlElement('w:i'))
    if bold:   rPr.append(OxmlElement('w:b'))
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t); link.append(run); paragraph._p.append(link)
    return link

def add_right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

def bottom_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1f4e3d')
    pbdr.append(b); pPr.append(pbdr)

# ---------------------------------------------------------------- running header / footer (every page)
_sec = doc.sections[0]
_hp = _sec.header.paragraphs[0]
_hp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
_set_run(_hp.add_run("\tSeunghyeon Lee  —  Curriculum Vitae"), 8, color=LIGHT)
_fp = _sec.footer.paragraphs[0]
_fp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
_set_run(_fp.add_run("\tSeunghyeon Lee    "), 8, color=LIGHT)
_pg = _fp.add_run()
for _t, _v in (("begin", None), ("instr", " PAGE "), ("end", None)):
    if _t == "instr":
        e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = _v
    else:
        e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), _t)
    _pg._r.append(e)
_set_run(_pg, 8, color=LIGHT)

# ---------------------------------------------------------------- header
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Seunghyeon Lee"); _set_run(r, 22, bold=True, color=INK)
r = p.add_run("  (Clay Lee)"); _set_run(r, 12, color=GREY)
r = p.add_run("    ·    Curriculum Vitae"); _set_run(r, 10.5, color=LIGHT, caps=True, spacing=20)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Ph.D., expected Aug 2026 · Forest Structure & Remote Sensing & GIS & AI")
_set_run(r, 10, color=ACCENT, bold=True)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Landscape & Ecological Planning Lab, Seoul National University, Seoul, Republic of Korea")
_set_run(r, 9.5, color=GREY)
p.paragraph_format.space_after = Pt(2)

contact = [
    ("shlee5598@snu.ac.kr", "mailto:shlee5598@snu.ac.kr"),
    ("+82-10-8455-5598", None),
    ("ORCID 0000-0003-4612-065X", "https://orcid.org/0000-0003-4612-065X"),
    ("Google Scholar", "https://scholar.google.com/citations?user=Ew1_-r0AAAAJ"),
    ("drseunghyeonlee.com", "https://drseunghyeonlee.com/en"),
]
p = doc.add_paragraph()
for i, (disp, url) in enumerate(contact):
    if i:
        sep = p.add_run("  |  "); _set_run(sep, 9, color=LIGHT)
    if url:
        add_hyperlink(p, url, disp, color="444444", size=9)
    else:
        _set_run(p.add_run(disp), 9, color=GREY)
bottom_border(p); p.paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------------- section helper
def section(title):
    p = doc.add_paragraph()
    r = p.add_run(title); _set_run(r, 11, bold=True, color=ACCENT, caps=True, spacing=30)
    bottom_border(p)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)

def entry(left_bold, right="", sub="", bullets=None, left_rest=""):
    """One CV entry: bold lead + optional right-aligned date, italic subline, bullets."""
    p = doc.add_paragraph(); add_right_tab(p)
    r = p.add_run(left_bold); _set_run(r, 9.5, bold=True)
    if left_rest:
        r = p.add_run(left_rest); _set_run(r, 9.5)
    if right:
        r = p.add_run("\t" + right); _set_run(r, 9, color=GREY)
    p.paragraph_format.space_before = Pt(5)
    if sub:
        ps = doc.add_paragraph(); r = ps.add_run(sub); _set_run(r, 9, italic=True, color=GREY)
    if bullets:
        for b in bullets:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent = Inches(0.25)
            pb.paragraph_format.space_after = Pt(0)
            r = pb.add_run(b); _set_run(r, 9)

def authors_run(p, authors, me="Seunghyeon Lee", mark_first=False):
    """Render byline; bold `me`. If mark_first and `me` is the first author, add † after the name."""
    for i, a in enumerate(authors):
        if i:
            s = p.add_run(", "); _set_run(s, 9, color=INK)
        r = p.add_run(a); _set_run(r, 9, bold=(a == me))
        if mark_first and i == 0 and a == me:
            d = p.add_run("†"); _set_run(d, 9, bold=True, color=ACCENT)

# ---------------------------------------------------------------- Research Interests
section("Research Interests")
# 첫 항목만 대문자로 시작하고 나머지는 이어쓰기 (한 문장으로 읽히게)
_ri = [r["Interest(EN)"] for r in sheet("Research Interests", order=lambda r: r["No."])]
p = doc.add_paragraph()
r = p.add_run(" · ".join([_ri[0]] + [_lower_first(x) for x in _ri[1:]]) + ".")
_set_run(r, 9.5)
p.paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- Education
section("Education")
def _edu_period(r):
    """수여 예정이면 CV 는 'Aug 2026' 형식으로 쓴다. 확정 학위는 Start-End 를 그대로."""
    se = str(r["Start–End"] or "")
    if not str(r["Date Degree Received"] or "").startswith("Expected"):
        return se
    head, _, tail = se.partition(" — ")
    y, _, m = tail.partition(".")
    return f"{head} — {_MONTHS[int(m)]} {y}" if m.isdigit() else se


def _edu_bullets(r):
    """학위논문 + (있으면) 심사위원. 박사는 Dissertation, 그 외는 Thesis."""
    out = []
    if r["학위논문(EN)"]:
        label = "Dissertation" if "Doctoral" in str(r["Level of Study"]) else "Thesis"
        out.append(f'{label}: “{r["학위논문(EN)"]}”')
    if r["심사위원(EN)"]:
        out.append(f'Committee: {r["심사위원(EN)"]}')
    return out or None


_MONTHS = {1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
           7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec"}

edu = [(f'{r["Degree(EN)"]}, {str(r["Program(EN)"]).split(" · ")[0]}',
        _edu_period(r), r["CV 서술(EN)"], _edu_bullets(r))
       for r in sheet("Education", where=lambda r: r["CV 표시"] == "Y",
                      order=lambda r: r["No."])]
for lead, date, sub, bl in edu:
    entry(lead, right=date, sub=sub, bullets=bl)

# ---------------------------------------------------------------- Research Experience
section("Research Experience")
exp = _work("research")
for lead, rest, date, sub in exp:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- Publications
section("Peer-Reviewed Publications")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  "
              "Impact Factor / quartile shown where available.  (SCI then KCI; newest first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

# --- source: citations/publications.bib (ascending) + IF/quartile from site.json (matched by DOI) ---
def _bib_entries(path):
    txt = open(path, encoding="utf-8").read()
    out, i = [], 0
    while True:
        at = txt.find("@", i)
        if at < 0: break
        b = txt.find("{", at); depth, j = 0, b
        while j < len(txt):
            if txt[j] == "{": depth += 1
            elif txt[j] == "}":
                depth -= 1
                if depth == 0: break
            j += 1
        out.append(txt[at:j + 1]); i = j + 1
    return out

def _bf(raw, name):
    m = re.search(r"\b" + name + r"\s*=\s*\{(.*?)\}\s*[,}\n]", raw, re.I | re.S)
    return m.group(1).strip() if m else ""

def _initials(given):
    return " ".join("-".join(s[0].upper() + "." for s in part.split("-") if s)
                    for part in given.split())

def _apa_author(name):
    if "," in name:
        sur, giv = [x.strip() for x in name.split(",", 1)]
    else:
        ps = name.split(); sur, giv = ps[-1], " ".join(ps[:-1])
    return f"{sur}, {_initials(giv)}".strip().rstrip(","), (sur == "Lee" and giv.startswith("Seunghyeon"))

_site = json.load(open(os.path.join(HERE, "..", "src", "data", "site.json"), encoding="utf-8"))
IFQ = {(it.get("doi") or "").lower().strip(): (it.get("if", ""), it.get("quartile", ""))
       for it in _site["publications"]["items"] if it.get("doi")}

_MO = {m: i for i, m in enumerate(
    ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"], 1)}
def _monthnum(raw):
    m = _bf(raw, "month").lower()[:3]
    return _MO.get(m, int(m) if m.isdigit() else 0)

_pubs = []
for raw in _bib_entries(os.path.join(HERE, "citations", "publications.bib")):
    doi = _bf(raw, "doi")
    iff, q = IFQ.get(doi.lower(), ("", ""))
    yv = _bf(raw, "year"); yv = int(yv) if yv.isdigit() else 0
    _pg = re.search(r"\d+", _bf(raw, "pages")); sp = int(_pg.group()) if _pg else 0
    _pubs.append({
        "apa": [_apa_author(a.strip()) for a in _bf(raw, "author").split(" and ")],
        "year": _bf(raw, "year"),
        "title": re.sub(r"\s*-\s*$", "", _bf(raw, "title")).replace(" - ", ": ").rstrip(". "),
        "journal": _bf(raw, "journal").replace("\\&", "&"),
        "vol": _bf(raw, "volume"), "num": _bf(raw, "number"),
        "pages": _bf(raw, "pages").replace("-", "–"), "doi": doi, "iff": iff, "q": q,
        "sortk": (0 if (iff or q) else 1, -yv, -_monthnum(raw), sp),  # SCI→KCI, newest first, then start page
    })
_pubs.sort(key=lambda d: d["sortk"])

for n, d in enumerate(_pubs, 1):
    apa, year, title, journal = d["apa"], d["year"], d["title"], d["journal"]
    vol, num, pages, doi, iff, q = d["vol"], d["num"], d["pages"], d["doi"], d["iff"], d["q"]
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({year}). "), 9)
    _set_run(p.add_run(title + ". "), 9)
    _set_run(p.add_run(journal), 9, italic=True)
    if vol:
        _set_run(p.add_run(", "), 9)
        _set_run(p.add_run(vol), 9, italic=True)
        if num:
            _set_run(p.add_run(f"({num})"), 9)
        if pages:
            _set_run(p.add_run(f", {pages}"), 9)
    _set_run(p.add_run(". "), 9)
    if doi:
        add_hyperlink(p, f"https://doi.org/{doi}", f"https://doi.org/{doi}", color="444444", size=9)
    if iff or q:
        iv = iff.split("(")[0].strip()
        tag = "  · " + ", ".join(x for x in [f"IF {iv}" if iv else "", q] if x)
        _set_run(p.add_run(tag), 8.5, color=ACCENT)

# ---------------------------------------------------------------- Manuscripts under review
section("Manuscripts Under Review")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  (5 manuscripts; first-author listed first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
# status = "stage|date"  (stage shown in brackets, date in grey)
# 심사중 원고: 저자는 쉼표 구분 문자열이므로 리스트로 되돌린다.
under_review = [([a.strip() for a in str(r["Authors(EN)"]).split(",")],
                 _dot(r["Title(EN)"]), r["Journal(EN)"],
                 f'{r["Status"]}|{r["Status Date"]}')
                for r in sheet("Under Review", order=lambda r: r["No."])]
for n, (auth, title, jour, status) in enumerate(under_review, 1):
    apa = [_apa_author(a) for a in auth]
    stage, _, dt = status.partition("|")
    yr = dt[:4] if dt[:4].isdigit() else "2026"
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({yr}). "), 9)
    _set_run(p.add_run(title.rstrip(". ")), 9, italic=True)       # APA: unpublished title in italics
    _set_run(p.add_run(f" [{stage}]. "), 9)
    _set_run(p.add_run(jour + "."), 9)
    if dt:
        _set_run(p.add_run(f"  · {dt}"), 8.5, color=GREY)

# ---------------------------------------------------------------- Books
section("Books")
# 저자 문자열 안에서 본인 이름만 굵게 — 이름 위치로 잘라 세 조각으로 출력한다
_ME = "Seunghyeon Lee"
for _n, _b in enumerate(sheet("Books", order=lambda r: r["No."]), 1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{_n}. "), 9, color=GREY)
    _auth = str(_b["Authors(EN)"] or "")
    _head, _sep, _tail = _auth.partition(_ME)
    _set_run(p.add_run(_head), 9)
    if _sep:
        _set_run(p.add_run(_sep), 9, bold=True)
    _set_run(p.add_run(f'{_tail} ({_b["Year"]}). '), 9)
    _set_run(p.add_run(str(_b["Title(EN)"])), 9, italic=True)
    _set_run(p.add_run(f'. {_b["Publisher(EN)"]}. ISBN {_b["ISBN"]}.'), 9)

# ---------------------------------------------------------------- Fellowships & Funding
section("Fellowships & Funding")
p = doc.add_paragraph()
r = p.add_run("Competitive fellowships / scholarships (amounts converted at ₩1,400 = $1).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
funding = [(r["CV 표기(EN)"], years_only(r["Period"]), r["Amount(USD)"])
           for r in sheet("Funding", order=lambda r: r["No."])]
for name, period, amount in funding:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
    _set_run(p.add_run(name), 9)
    _set_run(p.add_run(f"  ({amount})"), 9, color=ACCENT)
    _set_run(p.add_run("\t" + period), 9, color=GREY)

# ---------------------------------------------------------------- Invited talks
section("Invited Talks")
# 초청강연은 두 시트에 나뉘어 있다: 전용 시트(대학 초청강의) + 학회 시트의 CV분류=invited
invited = [(_dot(r["Title(EN)"]), r["Venue(EN)"], str(r["Year"]))
           for r in sheet("Invited Talks", order=lambda r: r["No."])] + \
          [(_dot(r["Title(EN)"]), f'{r["Conference(EN)"]}, {r["Venue(EN)"]}', str(r["Year"]))
           for r in sheet("Conferences", where=lambda r: r["CV 분류"] == "invited",
                          order=lambda r: _desc(r["Date"]))]
# 전 섹션 공통으로 최신순. 대학 초청강의(전용 시트)는 날짜가 없어 No. 순이며 모두 2025년이라
# 학회 시트에서 오는 2024년 건들보다 앞선다.
for title, venue, yr in invited:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    b = p.add_run("• "); _set_run(b, 9, color=ACCENT)
    r = p.add_run("Seunghyeon Lee"); _set_run(r, 9, bold=True)
    r = p.add_run(f" ({yr}). "); _set_run(r, 9)
    r = p.add_run(title + " "); _set_run(r, 9)
    r = p.add_run(venue + "."); _set_run(r, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Conferences
section("Conference Presentations (International)")
p = doc.add_paragraph()
r = p.add_run("First author / presenter unless marked “team”.  "
              "10 domestic presentations (KOSERT, KSEE, KILA) available on request.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

# 국제학회 발표만 CV 에 싣는다 (국내 10건은 요약 문구로 대체).
# 제목은 실제 초록 제출 제목을 저장하고, 표기만 APA 문장형으로 변환한다.
confs = [([a.strip() for a in str(r["Authors"]).split(",")], r["Year"],
          str(r["Type"]).capitalize() + (", team" if r["팀 발표"] == "Y" else ""),
          _dot(sentence_case(r["Title(EN)"])),
          f'{r["CV 학회(EN)"]}, {r["Venue(EN)"]}')
         for r in sheet("Conferences",
                        where=lambda r: r["Scope"] == "international"
                        and r["CV 분류"] == "conference",
                        order=lambda r: _desc(r["Date"]))]
for auth, yr, kind, title, venue in confs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
    authors_run(p, auth)
    yr_r = p.add_run(f" ({yr}). "); _set_run(yr_r, 9)
    t = p.add_run(title + " "); _set_run(t, 9)
    v = p.add_run(f"[{kind}] {venue}."); _set_run(v, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Grants / Projects
section("Research Projects")
p = doc.add_paragraph()
r = p.add_run("National R&D and municipal projects (newest first).  "
              "Budget = total project funding (national R&D).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

# 과제는 CV 표기(다듬은 과제명 · 기여 불릿 · 성과 요약)를 쓴다.
# xlsx 의 Title(EN)/기여(EN)/논문(EN)/특허(EN) 은 지원서 서식용 원문이라 별도로 둔다.
projects = [(r["CV 과제명(EN)"], r["Institute(EN)"],
             str(r["기간"] or "").replace("/", "."), r["Role(EN)"], r["PI(EN)"],
             r["예산(USD)"],
             str(r["CV 기여(EN)"]).splitlines() if r["CV 기여(EN)"] else None,
             r["CV 성과(EN)"])
            for r in sheet("Projects", order=lambda r: r["No."])]
for title, funder, period, role, pi, budget, bullets, outputs in projects:
    # 제목이 매우 길어 우측탭 기간이 잘리던 문제 → 제목은 전체 폭, 기간은 상세줄 맨 앞으로 이동
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    _set_run(p.add_run(title), 9.5, bold=True)
    ps = doc.add_paragraph()
    _set_run(ps.add_run(f"{period}"), 9, bold=True, color=GREY)
    _set_run(ps.add_run(" · "), 9, color=GREY)
    _set_run(ps.add_run(f"{funder} · "), 9, italic=True, color=GREY)
    _set_run(ps.add_run(f"Role: {role}"), 9, italic=True, bold=True, color=ACCENT)
    _set_run(ps.add_run(f" · PI: {pi}"), 9, italic=True, color=GREY)
    if budget:
        _set_run(ps.add_run(f" · {budget} (total project)"), 9, italic=True, color=GREY)
    for b in bullets:
        pb = doc.add_paragraph(style="List Bullet")
        pb.paragraph_format.left_indent = Inches(0.25); pb.paragraph_format.space_after = Pt(0)
        _set_run(pb.add_run(b), 9)
    if outputs:
        po = doc.add_paragraph(); po.paragraph_format.left_indent = Inches(0.25)
        po.paragraph_format.space_before = Pt(1)
        _set_run(po.add_run("→ Related outputs:  "), 8.5, bold=True, color=ACCENT)
        _set_run(po.add_run(outputs), 8.5, italic=True, color=GREY)

# ---------------------------------------------------------------- Patents
section("Patents (Republic of Korea)")
p = doc.add_paragraph()
r = p.add_run("7 registered, 4 pending.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

def _patent_sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _patent_list(items):
    for title, meta in items:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
        t = p.add_run(title + " "); _set_run(t, 9)
        m = p.add_run(meta); _set_run(m, 8.5, italic=True, color=GREY)

def _patents(status, datecol, fmt):
    """등록/출원 목록. 상세줄은 상태 + 날짜 + 권리자 규칙으로 만든다.
    제목은 공식 등록명이 아니라 CV 용으로 다듬은 표기를 쓴다."""
    out = []
    for r in sheet("Patents", where=lambda r: r["Status"] == status,
                   order=lambda r: _desc(r[datecol])):
        d = str(r[datecol] or "").replace("/", ".")
        out.append((r["CV 특허명(EN)"], f'{fmt} {d[:7] if fmt == "Registered" else d[:4]}'
                    f' · {r["CV 권리자(EN)"]}'))
    return out


_patent_sub("Registered")
_patent_list(_patents("registered", "Reg Date", "Registered"))
_patent_sub("Pending")
_patent_list(_patents("application", "App Date", "Pending"))

# ---------------------------------------------------------------- Honors
section("Awards & Honors")

def _subhead(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _tabbed(rows):
    for h, d in rows:
        p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
        r = p.add_run(h); _set_run(r, 9)
        r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

_subhead("Awards")
_tabbed([(f'{r["Award(EN)"]} — {r["Organization(EN)"]} '
          f'({_lower_first(r["Details(EN)"])})', str(r["Year"]))
         for r in sheet("Awards", order=lambda r: _desc(r["Year"]))])

# ---------------------------------------------------------------- Service & Membership
section("Professional Service & Membership")
for _s in sheet("Service & Membership", where=lambda r: r["Type"] == "service",
                order=lambda r: r["No."]):
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
    _set_run(p.add_run(_s["Role(EN)"] + ", "), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(_s["Organization(EN)"]), 9)
    _set_run(p.add_run("\t" + _s["Period"]), 9, color=GREY)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
_set_run(p.add_run("Professional memberships"), 9, bold=True, color=ACCENT)
members = [(r["Organization(EN)"], r["Period"])
           for r in sheet("Service & Membership", where=lambda r: r["Type"] == "membership",
                          order=lambda r: r["No."])]
for name, period in members:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(3)
    _set_run(p.add_run("Member, " + name), 9)
    _set_run(p.add_run("\t" + period), 9, color=GREY)

# ---------------------------------------------------------------- Teaching
section("Teaching Experience")
_tw = _site["teaching"]["items"]
_n_univ   = sum(1 for _t in _tw if _t["type"] == "univ")
_n_online = sum(1 for _t in _tw if _t["type"] == "online")
_meta     = _site.get("meta", {})
_students = _meta.get("onlineStudents", "1,100+")
_sessions = _meta.get("invitedSessions", "20+")
_years    = _meta.get("invitedYears", "2022 — 2026")
p = doc.add_paragraph()
r = p.add_run(f"{_n_univ} university courses · {_n_online} online courses ({_students} students) · {_sessions} invited lectures and workshops.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

def _t_sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _t_item(text, date, url=None):
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    b = p.add_run("• "); _set_run(b, 9, color=ACCENT)
    if url:
        add_hyperlink(p, url, text, color="1a1a1a", size=9)   # clickable title, normal colour
    else:
        r = p.add_run(text); _set_run(r, 9)
    r = p.add_run("\t" + date); _set_run(r, 9, color=GREY)

# CV 표기(EN) 가 채워진 행만 싣는다. 여러 회차를 한 줄로 묶은 경우 대표 행에만 표기가 있다.
def _teach(kind):
    return sheet("Teaching", where=lambda r: r["Type"] == kind and r["CV 표기(EN)"],
                 order=lambda r: _desc(r["Start"]))


_t_sub("University Courses")
for _u in _teach("univ"):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run("• "), 9, color=ACCENT)
    _set_run(p.add_run(_u["CV 표기(EN)"]), 9)
    _lines = str(_u["CV 서술(EN)"] or "").splitlines()
    for _i, _line in enumerate(_lines):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.28)
        if _i == len(_lines) - 1:
            add_right_tab(p)
        _set_run(p.add_run(_line), 9, color=GREY)
        if _i == len(_lines) - 1:
            _set_run(p.add_run("	" + str(_u["CV 기간(EN)"])), 9, color=GREY)

_t_sub(f"Online Courses  ({_students} students enrolled)")
for _o in _teach("online"):
    _t_item(_o["CV 표기(EN)"], str(_o["CV 기간(EN)"]), _o["URL"])

_t_sub(f"Invited Lectures & Workshops  ({_sessions} sessions, {_years})")
for _v in _teach("special"):
    _t_item(_v["CV 표기(EN)"], str(_v["CV 기간(EN)"]))

# ---------------------------------------------------------------- Professional Experience
section("Professional Experience")
prof = _work("professional")
for lead, rest, date, sub in prof:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- References
section("References")
# 공개 CV에는 추천인 연락처를 노출하지 않음(사생활). 실제 3인 연락처는 비공개 레퍼런스 시트로 제출.
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
_set_run(p.add_run("Available upon request."), 9, italic=True, color=GREY)

# (Technical Skills section intentionally omitted)

# ---------------------------------------------------------------- date / signature
# Drop a signature image at cv/assets/signature.png (transparent PNG, tightly cropped) to embed it;
# otherwise a blank signature line is printed.
_sig_img = os.path.join(HERE, "assets", "signature.png")
_today = _dtdate.today().strftime("%d %B %Y")
if os.path.exists(_sig_img):
    pim = doc.add_paragraph(); pim.paragraph_format.space_before = Pt(16)
    pim.add_run().add_picture(_sig_img, height=Inches(0.6))
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(1)
    _set_run(p.add_run("Seunghyeon Lee"), 9.5, bold=True)
    _set_run(p.add_run("\tDate: " + _today), 9, color=GREY)
else:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(18)
    _set_run(p.add_run("Seunghyeon Lee"), 9.5, bold=True)
    _set_run(p.add_run("\tSignature: ____________________     Date: " + _today), 9, color=GREY)

# ---------------------------------------------------------------- save
out = "Lee_Seunghyeon_CV_v11.docx"
doc.save(out)
print("saved:", out)
