# -*- coding: utf-8 -*-
"""
Academic CV builder — Seunghyeon (Clay) Lee
Forest remote sensing / LiDAR · GEDI / ecology — postdoc application

Structure follows field-standard academic CVs (Pascual UMD / Harvard GSAS).
All content is sourced from the verified portfolio data in
  ../src/data/site.json
Edit the data blocks below, then run:  python build_cv.py
Output: Lee_Seunghyeon_CV_v1.docx  ->  export to PDF in Word for ../public/cv/
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, json

HERE = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------- palette
INK    = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x1f, 0x4e, 0x3d)   # deep forest green
GREY   = RGBColor(0x55, 0x55, 0x55)
LIGHT  = RGBColor(0x88, 0x88, 0x88)
BODYFONT = "Calibri"
HEADFONT = "Calibri"

doc = Document()

# page margins
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

RIGHT_TAB = Inches(6.9)  # within 7.0" text width

# base style
st = doc.styles["Normal"]
st.font.name = BODYFONT
st.font.size = Pt(9.5)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.04

def _set_run(r, size=9.5, bold=False, italic=False, color=INK, font=BODYFONT, caps=False, spacing=None):
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps: r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr(); sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing)); rPr.append(sp)

def add_right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

def bottom_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1f4e3d')
    pbdr.append(b); pPr.append(pbdr)

# ---------------------------------------------------------------- header
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Seunghyeon Lee"); _set_run(r, 22, bold=True, color=INK)
r = p.add_run("  (이승현 · Clay Lee)"); _set_run(r, 12, color=GREY)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Ph.D. Candidate (expected Aug 2026) · Forest Structure & Remote Sensing (LiDAR / GEDI)")
_set_run(r, 10, color=ACCENT, bold=True)
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
r = p.add_run("Landscape & Ecological Planning Lab, Seoul National University, Seoul, Republic of Korea")
_set_run(r, 9.5, color=GREY)
p.paragraph_format.space_after = Pt(2)

contact = [
    "shlee5598@snu.ac.kr",
    "+82-10-8455-5598",
    "shleeclay.github.io",
    "github.com/shleeclay",
    "ORCID: 0000-0003-4612-065X",
]
p = doc.add_paragraph()
for i, c in enumerate(contact):
    if i:
        sep = p.add_run("   |   "); _set_run(sep, 9, color=LIGHT)
    r = p.add_run(c); _set_run(r, 9, color=GREY)
bottom_border(p); p.paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------------- section helper
def section(title):
    p = doc.add_paragraph()
    r = p.add_run(title); _set_run(r, 11, bold=True, color=ACCENT, caps=True, spacing=30)
    bottom_border(p)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)

def entry(left_bold, right="", sub="", bullets=None, left_rest=""):
    """One CV entry: bold lead + optional right-aligned date, italic subline, bullets."""
    p = doc.add_paragraph(); add_right_tab(p)
    r = p.add_run(left_bold); _set_run(r, 9.5, bold=True)
    if left_rest:
        r = p.add_run(left_rest); _set_run(r, 9.5)
    if right:
        r = p.add_run("\t" + right); _set_run(r, 9, color=GREY)
    p.paragraph_format.space_before = Pt(3)
    if sub:
        ps = doc.add_paragraph(); r = ps.add_run(sub); _set_run(r, 9, italic=True, color=GREY)
    if bullets:
        for b in bullets:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent = Inches(0.25)
            pb.paragraph_format.space_after = Pt(0)
            r = pb.add_run(b); _set_run(r, 9)

def authors_run(p, authors, me="Seunghyeon Lee", mark_first=False):
    """Render byline; bold `me`. If mark_first and `me` is the first author, add † after the name."""
    for i, a in enumerate(authors):
        if i:
            s = p.add_run(", "); _set_run(s, 9, color=INK)
        r = p.add_run(a); _set_run(r, 9, bold=(a == me))
        if mark_first and i == 0 and a == me:
            d = p.add_run("†"); _set_run(d, 9, bold=True, color=ACCENT)

# ---------------------------------------------------------------- Research Interests
section("Research Interests")
p = doc.add_paragraph()
r = p.add_run("Airborne & spaceborne LiDAR (GEDI) for vertical forest structure · forest typology and "
              "stratification metrics (FHD, PAI, layer indices) · biomass and carbon estimation · "
              "urban ecosystems and green infrastructure · UAV/thermal wildlife monitoring · "
              "reproducible, code-driven remote-sensing science.")
_set_run(r, 9.5)
p.paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- Education
section("Education")
edu = [
    ("Ph.D. Candidate, Interdisciplinary Program in Landscape Architecture",
     "2022.03 — Aug 2026 (expected)",
     "Integrated Major in SmartCity Global Convergence · Seoul National University · "
     "Landscape & Ecological Planning Lab · Advisor: Prof. Youngkeun Song",
     ["Dissertation: “LiDAR-Based Characterization of Vertical Structure and Disturbance in "
      "Temperate Forests”",
      "Committee: Dongkun Lee, Youngryel Ryu (Seoul National Univ.); Hansoo Kim (Gyeonggi Research "
      "Institute); Dennis Heejoon Choi (Dankook Univ.); Youngkeun Song (advisor, SNU)"]),
    ("M.A., Landscape Architecture",
     "2020.03 — 2022.02",
     "Integrated Major in SmartCity Global Convergence · Graduate School of Environmental Studies, "
     "Seoul National University · Advisor: Prof. Youngkeun Song",
     ["Thesis: “Feasibility of Wildlife Detection Using UAV-derived Thermal and True-color Imagery”"]),
    ("B.S., Landscape Architecture",
     "2011.03 — 2017.02",
     "College of Agriculture & Life Sciences, Kyungpook National University, Daegu, Republic of Korea", None),
]
for lead, date, sub, bl in edu:
    entry(lead, right=date, sub=sub, bullets=bl)

# ---------------------------------------------------------------- Research Experience
section("Research Experience")
exp = [
    ("Visiting Research Intern", " — Purdue University, FNR FUSE Lab (PI: Prof. Brady Hardiman)",
     "2023.12 — 2024.02", "West Lafayette, IN, United States · Forest structure & LiDAR"),
    ("Visiting Research Intern", " — Virginia Tech (PI: Prof. Jaeyoung Ha)",
     "2024.07 — 2024.08", "Blacksburg, VA, United States · Landscape architecture"),
    ("Graduate Researcher", " — Landscape & Ecological Planning Lab, Seoul National University",
     "2020.03 — Present", "National R&D and municipal projects on LiDAR forest structure, carbon, and urban ecology"),
    ("Research Assistant", " — Architecture & Urban Research Institute (AURI)",
     "2019.08 — 2019.11", "Republic of Korea"),
]
for lead, rest, date, sub in exp:
    entry(lead, right=date, sub=sub, left_rest=rest)

# ---------------------------------------------------------------- Publications
section("Peer-Reviewed Publications")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  "
              "Impact Factor / quartile shown where available.  (SCI then KCI; newest first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

# --- source: citations/publications.bib (ascending) + IF/quartile from site.json (matched by DOI) ---
def _bib_entries(path):
    txt = open(path, encoding="utf-8").read()
    out, i = [], 0
    while True:
        at = txt.find("@", i)
        if at < 0: break
        b = txt.find("{", at); depth, j = 0, b
        while j < len(txt):
            if txt[j] == "{": depth += 1
            elif txt[j] == "}":
                depth -= 1
                if depth == 0: break
            j += 1
        out.append(txt[at:j + 1]); i = j + 1
    return out

def _bf(raw, name):
    m = re.search(r"\b" + name + r"\s*=\s*\{(.*?)\}\s*[,}\n]", raw, re.I | re.S)
    return m.group(1).strip() if m else ""

def _initials(given):
    return " ".join("-".join(s[0].upper() + "." for s in part.split("-") if s)
                    for part in given.split())

def _apa_author(name):
    if "," in name:
        sur, giv = [x.strip() for x in name.split(",", 1)]
    else:
        ps = name.split(); sur, giv = ps[-1], " ".join(ps[:-1])
    return f"{sur}, {_initials(giv)}".strip().rstrip(","), (sur == "Lee" and giv.startswith("Seunghyeon"))

_site = json.load(open(os.path.join(HERE, "..", "src", "data", "site.json"), encoding="utf-8"))
IFQ = {(it.get("doi") or "").lower().strip(): (it.get("if", ""), it.get("quartile", ""))
       for it in _site["publications"]["items"] if it.get("doi")}

_MO = {m: i for i, m in enumerate(
    ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"], 1)}
def _monthnum(raw):
    m = _bf(raw, "month").lower()[:3]
    return _MO.get(m, int(m) if m.isdigit() else 0)

_pubs = []
for raw in _bib_entries(os.path.join(HERE, "citations", "publications.bib")):
    doi = _bf(raw, "doi")
    iff, q = IFQ.get(doi.lower(), ("", ""))
    yv = _bf(raw, "year"); yv = int(yv) if yv.isdigit() else 0
    _pg = re.search(r"\d+", _bf(raw, "pages")); sp = int(_pg.group()) if _pg else 0
    _pubs.append({
        "apa": [_apa_author(a.strip()) for a in _bf(raw, "author").split(" and ")],
        "year": _bf(raw, "year"),
        "title": re.sub(r"\s*-\s*$", "", _bf(raw, "title")).replace(" - ", ": ").rstrip(". "),
        "journal": _bf(raw, "journal").replace("\\&", "&"),
        "vol": _bf(raw, "volume"), "num": _bf(raw, "number"),
        "pages": _bf(raw, "pages").replace("-", "–"), "doi": doi, "iff": iff, "q": q,
        "sortk": (0 if (iff or q) else 1, -yv, -_monthnum(raw), sp),  # SCI→KCI, newest first, then start page
    })
_pubs.sort(key=lambda d: d["sortk"])

for n, d in enumerate(_pubs, 1):
    apa, year, title, journal = d["apa"], d["year"], d["title"], d["journal"]
    vol, num, pages, doi, iff, q = d["vol"], d["num"], d["pages"], d["doi"], d["iff"], d["q"]
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({year}). "), 9)
    _set_run(p.add_run(title + ". "), 9)
    _set_run(p.add_run(journal), 9, italic=True)
    if vol:
        _set_run(p.add_run(", "), 9)
        _set_run(p.add_run(vol), 9, italic=True)
        if num:
            _set_run(p.add_run(f"({num})"), 9)
        if pages:
            _set_run(p.add_run(f", {pages}"), 9)
    _set_run(p.add_run(". "), 9)
    if doi:
        _set_run(p.add_run(f"https://doi.org/{doi}"), 9, color=GREY)
    if iff or q:
        iv = iff.split("(")[0].strip()
        tag = "  · " + ", ".join(x for x in [f"IF {iv}" if iv else "", q] if x)
        _set_run(p.add_run(tag), 8.5, color=ACCENT)

# ---------------------------------------------------------------- Manuscripts under review
section("Manuscripts Under Review")
p = doc.add_paragraph()
r = p.add_run("APA 7th.  Name in bold = author; † = first author.  (5 manuscripts; first-author listed first)")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)
under_review = [
    (["Seunghyeon Lee", "Dennis Heejoon Choi", "Youngkeun Song", "James H. Thorne"],
     "Spaceborne LiDAR reveals post-fire vertical structural loss in a dense temperate Asian conifer "
     "plantation.", "Science of Remote Sensing", "Under review (2026.04)"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"],
     "Systematic bias in urban-forest vegetation coverage assessment: the influence of topo-edaphic "
     "conditions on discrepancies between ALS and field surveys.", "Geoscience Letters",
     "In review (2026.06)"),
    (["Seunghyeon Lee", "Yonghwan Kim", "Dohee Kim", "Hansoo Kim", "Youngkeun Song"],
     "Bi-temporal ALS assessment of vertical–horizontal forest structures and their structural "
     "association in a temperate urban forest.", "Forest Science and Technology", "In review (2026.05)"),
    (["Yonghwan Kim", "Seunghyeon Lee", "Wonhyeop Shin", "Youngkeun Song"],
     "Integrating UAV-derived habitat metrics with movement persistence and species distribution models "
     "to characterize seasonal habitat use of invasive turtles in urban wetlands.",
     "Global Ecology and Conservation", "In review (2026.05)"),
    (["Gapseong Jekal", "Yong Hwan Kim", "Seunghyeon Lee", "Ji Weon Yun", "Dae Yeol Kim", "Youngkeun Song"],
     "Monitoring transition-zone dynamics of a Phragmites communis–Suaeda japonica mosaic from "
     "multi-season Sentinel-2 in a coastal wetland.", "Journal of Coastal Conservation",
     "In review (2026.02)"),
]
for n, (auth, title, jour, status) in enumerate(under_review, 1):
    apa = [_apa_author(a) for a in auth]
    ym = re.search(r"\((20\d\d(?:\.\d+)?)\)", status)
    yr = (ym.group(1).split(".")[0] if ym else "2026")
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    _set_run(p.add_run(f"{n}. "), 9, color=GREY)
    nA = len(apa)
    for i, (disp, is_user) in enumerate(apa):
        if i > 0:
            _set_run(p.add_run(", "), 9)
        if i == nA - 1 and nA > 1:
            _set_run(p.add_run("& "), 9)
        _set_run(p.add_run(disp), 9, bold=is_user)
        if i == 0 and is_user:
            _set_run(p.add_run("†"), 9, bold=True, color=ACCENT)
    _set_run(p.add_run(f" ({yr}). "), 9)
    _set_run(p.add_run(title.rstrip(". ")), 9, italic=True)       # APA: unpublished title in italics
    _set_run(p.add_run(" [Manuscript under review]. "), 9)
    _set_run(p.add_run(jour + "."), 9)
    if ym:
        _set_run(p.add_run(f"  · submitted {ym.group(1)}"), 8.5, color=GREY)

# ---------------------------------------------------------------- Invited talks
section("Invited Talks")
invited = [
    ("Remotely sensed smart-city ecological value maps.",
     "NEF (Neocity Empowerment Forum), Orlando, FL, US", "2024"),
    ("Invasive species monitoring development and its application.",
     "Asia Week, Fukuoka, Japan", "2024"),
]
for title, venue, yr in invited:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28)
    b = p.add_run("• "); _set_run(b, 9, color=ACCENT)
    r = p.add_run("Seunghyeon Lee"); _set_run(r, 9, bold=True)
    r = p.add_run(f" ({yr}). "); _set_run(r, 9)
    r = p.add_run(title + " "); _set_run(r, 9)
    r = p.add_run(venue + "."); _set_run(r, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Conferences
section("Conference Presentations (International)")
p = doc.add_paragraph()
r = p.add_run("First author / presenter unless marked “team”.  "
              "10 domestic presentations (KOSERT, KSEE, KILA) available on request.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

confs = [
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2025, "Poster",
     "Novel LiDAR indices reveal scale-dependent vertical structure and typologies in a temperate forest.",
     "AGU Fall Meeting, New Orleans, US"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Poster",
     "Quantifying the number of forest stratification layers of city-scale forest and characterizing by "
     "forest types.", "AGU Fall Meeting, Washington D.C., US"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Oral",
     "Canopy layers distribution by forest patch and species at the city-level forest using airborne LiDAR.",
     "ICLEE, Kitakyushu, Japan"),
    (["Seunghyeon Lee", "Hansoo Kim", "Youngkeun Song"], 2024, "Poster",
     "Characterizing forest vegetation stratification by forest biotope types using airborne LiDAR.",
     "ESA Annual Meeting, Long Beach, CA, US"),
    (["BK21 SmartCity Team"], 2024, "Video, team",
     "Smart city & green infrastructure.", "CES (Consumer Electronics Show), Las Vegas, NV, US"),
    (["Landscape & Ecological Planning Lab"], 2023, "Oral, team",
     "Property-information analysis of Gyeonggi-do Province biotopes.",
     "JCK Symposium (18th), Kyoto, Japan"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2023, "Poster",
     "Wildfire-driven forest vegetation height change comparison.", "JCK Symposium (18th), Kyoto, Japan"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2023, "Poster",
     "The impacts of fire-induced disturbances on tree height and structure using GEDI-derived variables.",
     "ESA Annual Meeting, Portland, OR, US"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Monitoring structural change of fire-induced forest vegetation using GEDI.", "ICLEE (12th), Online"),
    (["Seunghyeon Lee", "Youngkeun Song"], 2022, "Oral",
     "Comparison of GEDI and aerial laser scanning datasets according to leaf-on and leaf-off seasons.",
     "AGU Fall Meeting, Chicago, IL, US / Online"),
    (["Seunghyeon Lee", "Dennis Heejoon Choi", "Youngkeun Song"], 2022, "Video",
     "Classification of tree species and forest floor using an XGBoost algorithm with forest maps, "
     "airborne LiDAR, and satellite imagery.", "World Forestry Congress (15th), Seoul, Republic of Korea"),
    (["Seunghyeon Lee", "Sung-Ho Kil", "Youngkeun Song"], 2022, "Poster",
     "Feasibility analyses of real-time detection of wildlife using UAV-derived thermal and RGB images.",
     "World Forestry Congress (15th), Seoul, Republic of Korea"),
    (["Seunghyeon Lee", "DaeYeol Kim", "Dennis Heejoon Choi", "Hansoo Kim", "Youngkeun Song"], 2021, "Poster",
     "Detecting individual broad-leaved trees by a trunk-extraction method using leaf-off airborne LiDAR.",
     "AGU Fall Meeting, New Orleans, US / Online"),
]
for auth, yr, kind, title, venue in confs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
    authors_run(p, auth)
    yr_r = p.add_run(f" ({yr}). "); _set_run(yr_r, 9)
    t = p.add_run(title + " "); _set_run(t, 9)
    v = p.add_run(f"[{kind}] {venue}."); _set_run(v, 9, italic=True, color=GREY)

# ---------------------------------------------------------------- Patents
section("Patents (Republic of Korea)")
p = doc.add_paragraph()
r = p.add_run("7 registered, 4 pending.")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_after = Pt(3)

def _patent_sub(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _patent_list(items):
    for title, meta in items:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        bullet = p.add_run("• "); _set_run(bullet, 9, color=ACCENT)
        t = p.add_run(title + " "); _set_run(t, 9)
        m = p.add_run(meta); _set_run(m, 8.5, italic=True, color=GREY)

_patent_sub("Registered")
_patent_list([
    ("System and method for providing a tree-management platform.", "Registered 2025.09 · Sole inventor"),
    ("Method and system for determining wetland vegetation structure using drone LiDAR.",
     "Registered 2025.04 · Seoul National University"),
    ("Image-based roadside-tree information management system and method.", "Registered 2024.11 · Sole inventor"),
    ("Prediction of missing location points of a wildlife GPS tracker by combining in-situ and biological "
     "information.", "Registered 2023.07 · Seoul National University"),
    ("Self-heating target module for thermal-camera performance testing and precision data acquisition.",
     "Registered 2022.10 · Seoul National University"),
    ("Automated wildlife detection method and device using drone-mounted thermal-camera imagery.",
     "Registered 2022.05 · Seoul National University"),
    ("A method to derive an optimal plan for outdoor thermal-comfort mitigation.",
     "Registered 2022.04 · Seoul National University"),
])
_patent_sub("Pending")
_patent_list([
    ("Quantitative assessment and automatic classification of wildfire damage using spaceborne LiDAR.",
     "Pending 2025 · Dankook Univ. & Seoul National University"),
    ("Monitoring halophyte distribution from multi-period vegetation indices in coastal wetlands.",
     "Pending 2025 · Seoul National University"),
    ("Determining terrestrialization of abandoned-paddy wetlands using multi-temporal drone LiDAR and "
     "stratigraphic volume change.", "Pending 2023 · Seoul National University"),
    ("Deriving optimal wildlife-capture range via machine learning by fusing GPS-tracking data and drone "
     "footage.", "Pending 2023 · Seoul National University"),
])

# ---------------------------------------------------------------- Grants / Projects
section("Research Projects & Grants")
projects = [
    ("Integrated assessment tool for carbon reduction by offset and synergy with ecosystem services",
     "Korean Ministry of Environment", "2023 — 2027", "Researcher"),
    ("Creation, restoration & management technology of carbon-accumulated abandoned-paddy wetland",
     "Korean Ministry of Environment", "2022 — 2026", "Researcher"),
    ("Carbon storage/sink analysis and inventory for Gyeonggi-do considering biotope types",
     "Gyeonggi Research Institute", "2024 — 2025", "Researcher"),
    ("Real-time web-based positioning surveillance system for introduced exotic species",
     "Korean Ministry of Environment", "2021 — 2023", "Project Manager"),
    ("Redundancy-based green-infrastructure technology for urban ecosystem challenges",
     "Korean Ministry of Environment", "2020 — 2022", "Researcher"),
]
for title, inst, period, role in projects:
    entry(title, right=period, sub=f"{inst} · {role}")
p = doc.add_paragraph(); r = p.add_run("…and 7 additional national/municipal R&D projects (2018 — present).")
_set_run(r, 8.5, italic=True, color=LIGHT); p.paragraph_format.space_before = Pt(2)

# ---------------------------------------------------------------- Honors
section("Awards & Honors")

def _subhead(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); _set_run(r, 9.5, bold=True, color=ACCENT)

def _tabbed(rows):
    for h, d in rows:
        p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h); _set_run(r, 9)
        r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

_subhead("Awards")
_tabbed([
    ("Best Presentation Award — KOSERT Autumn Meeting "
     "(GEDI vs. airborne LiDAR field agreement by vegetation type and season)", "2022"),
    ("Best Presentation Award — KOSERT Spring Meeting "
     "(forest layer-structure indices and their correlation with carbon storage)", "2025"),
])
_subhead("Scholarships & Fellowships")
_tabbed([
    ("Chung Mong-Koo Foundation — OnDream Future Industry Talent Graduate Scholarship (Full)", "2022 — Present"),
    ("BK21 FOUR — Integrated Major in SmartCity Global Convergence (Full)", "2020 — Present"),
    ("Ilju Foundation Undergraduate Scholarship, 23rd Scholar (Full)", "2015 — 2017"),
    ("Challenge Scholarship, Kyungpook National University (Full)", "2011 — 2013"),
])

# ---------------------------------------------------------------- Service & Membership
section("Professional Service & Membership")
p = doc.add_paragraph()
r = p.add_run("Peer reviewer:  "); _set_run(r, 9, bold=True, color=ACCENT)
r = p.add_run("International Journal of Applied Earth Observation and Geoinformation (2026 – present)")
_set_run(r, 9)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
r = p.add_run("Member:  "); _set_run(r, 9, bold=True, color=ACCENT)
r = p.add_run("American Geophysical Union (AGU, 2021–); Ecological Society of America (ESA, 2021–); "
              "Korean Society of Environmental Restoration Technology (KOSERT, 2020–); "
              "Ecological Society of Korea (2023–)")
_set_run(r, 9)

# ---------------------------------------------------------------- Teaching
section("Teaching & Outreach")
teach = [
    ("University Lecturer (sole instructor) — “Understanding and Application of Spatial Information” "
     "(3 cr.; ~30 students/yr), Incheon National University — GIS theory & QGIS labs, "
     "remote-sensing theory & analysis", "2024, 2025"),
    ("Online Instructor — 4 GIS/QGIS courses on an online learning platform "
     "(1,000+ enrolled students total)", "2022 — 2025"),
    ("Invited Workshops — K-water, Korea Environment Corporation, Hyundai NGV, SeSAC, SNU (20+ sessions)",
     "2020 — 2025"),
]
for t, d in teach:
    p = doc.add_paragraph(); add_right_tab(p); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(t); _set_run(r, 9)
    r = p.add_run("\t" + d); _set_run(r, 9, color=GREY)

# ---------------------------------------------------------------- Skills
section("Technical Skills & Certifications")
skills = [
    ("Programming & analysis", "Python, R, Google Earth Engine"),
    ("LiDAR / point cloud", "ALS · GEDI · MLS/TLS processing, LAStools / PDAL, canopy & vertical-structure metrics"),
    ("Remote sensing & GIS", "QGIS, ArcGIS, satellite & UAV imagery, thermal imaging, Envi-met microclimate"),
    ("Machine learning", "classification / regression (XGBoost, deep learning) for forest & wildlife mapping"),
    ("Certifications", "Engineer Forest, Engineer Landscape Architecture, Drone Pilot License (Class 2)"),
    ("Languages", "Korean (native), English (professional working), Chinese (elementary)"),
]
for label, val in skills:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(label + ":  "); _set_run(r, 9, bold=True, color=ACCENT)
    r = p.add_run(val); _set_run(r, 9)

# ---------------------------------------------------------------- References
section("References")
refs = [
    ("Prof. Youngkeun Song (Doctoral advisor)",
     "Landscape & Ecological Planning Lab, Seoul National University · [email — to be added]"),
    ("Prof. Brady S. Hardiman",
     "Dept. of Forestry & Natural Resources, Purdue University · [email — to be added]"),
    ("Prof. Jaeyoung Ha",
     "School of Architecture + Design, Virginia Tech · [email — to be added]"),
]
for name, detail in refs:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(name + " — "); _set_run(r, 9, bold=True)
    r = p.add_run(detail); _set_run(r, 9, color=GREY)

# ---------------------------------------------------------------- save
out = "Lee_Seunghyeon_CV_v2.docx"
doc.save(out)
print("saved:", out)
