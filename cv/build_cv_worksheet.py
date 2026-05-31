# -*- coding: utf-8 -*-
"""
CV fill-in worksheet (combined) — Parts 1~6 in one file.
Pre-fills what is known from site.json; leaves red-underlined blanks to complete.
Run:  python build_cv_worksheet.py  ->  CV_worksheet.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK    = RGBColor(0x1a, 0x1a, 0x1a)
ACCENT = RGBColor(0x1f, 0x4e, 0x3d)
GREY   = RGBColor(0x55, 0x55, 0x55)
BLANK  = RGBColor(0xb0, 0x00, 0x20)
LINE   = "________________________________"

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.05

def run(p, text, size=10, bold=False, italic=False, color=INK):
    r = p.add_run(text); r.font.name = "Calibri"; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color; return r

def para(space_before=0, space_after=0):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after); return p

def bottom_border(p, color='1f4e3d'):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),color)
    pbdr.append(b); pPr.append(pbdr)

def section(title, note=""):
    p = para(16, 2); run(p, title, 13, bold=True, color=ACCENT); bottom_border(p)
    if note:
        p = para(0, 4); run(p, note, 9, italic=True, color=GREY)

def field(label, hint="", indent=0.15):
    p = para(2, 0); p.paragraph_format.left_indent = Inches(indent)
    run(p, label + ": ", 10, bold=True, color=BLANK)
    run(p, LINE + ("  " + hint if hint else ""), 10, color=BLANK)

def fixed(label, value):
    p = para(2, 0); p.paragraph_format.left_indent = Inches(0.15)
    run(p, label + ": ", 10, bold=True); run(p, value, 10, color=GREY)

def subhead(text):
    p = para(8, 1); p.paragraph_format.left_indent = Inches(0.1)
    run(p, text, 10.5, bold=True)

def blank_block(fields, count, indent=0.3):
    for i in range(1, count + 1):
        p = para(4, 0); p.paragraph_format.left_indent = Inches(indent)
        run(p, f"({i})", 10, bold=True, color=BLANK)
        for label, hint in fields:
            field("   " + label, hint, indent=indent)

# ============================= cover =============================
p = para(0, 2); run(p, "CV 작성 워크시트 — Parts 1~6", 18, bold=True, color=ACCENT); bottom_border(p)
p = para(0, 2)
run(p, "검정 = 확정(수정 자유) · ", 9, color=GREY)
run(p, "빨강 = 직접 채워주실 칸", 9, color=BLANK, bold=True)
p = para(0, 2)
run(p, "해당 없는 항목은 비워두시면 CV에서 자동 제외합니다. 핵심 성과(논문·특허)와 연결되는 항목 위주로 채우셔도 충분합니다.",
    9, italic=True, color=GREY)
for ln in ["Part 1 — 학위·논문", "Part 2 — 출판 파이프라인", "Part 3 — 연구과제",
           "Part 4 — 수상·명예", "Part 5 — 학계 기여", "Part 6 — 교육·멘토링"]:
    p = para(0, 0); p.paragraph_format.left_indent = Inches(0.2); run(p, "•  " + ln, 9.5, color=GREY)

# ============================= PART 1 =============================
section("Part 1 — 학위 · 논문 (Education details)",
        "학위논문 제목과 심사위원은 두 예시 CV 모두 명시한 항목입니다.")
fixed("PhD 지도교수", "Prof. Youngkeun Song (송영근) — 확정")
field("PhD dissertation 영문 제목")
field("PhD 심사위원(지도교수 외) 명단", "(이름 · 소속)")
field("MA thesis 영문 제목")
field("MA 심사위원(선택)", "(이름)")

# ============================= PART 2 =============================
section("Part 2 — 출판 파이프라인 (Manuscripts & Metrics)",
        "포닥 심사에서 '진행중 연구력'을 보여주는 핵심. 게재 9편은 이미 CV에 있음.")
subhead("◦ Under review (투고·심사중)")
blank_block([("저자", "(이승현 위치 표시)"), ("제목", ""),
             ("대상 저널", ""), ("상태·투고일", "(예: Under review, 2026.03)")], count=3)
subhead("◦ In preparation (작성중)")
blank_block([("저자", ""), ("제목", ""), ("대상 저널", "")], count=3)
subhead("◦ 지표·링크")
field("Google Scholar URL")
field("ORCID URL")
field("총 인용수 / h-index", "(선택 — Google Scholar 기준)")

# ============================= PART 3 (projects) =============================
section("Part 3 — 연구과제 (Research Projects)",
        "과제마다 ①연구비 ②과제책임자(PI) ③본인이 한 일 2~4불릿 을 채워주세요. "
        "산출물은 제안이며 맞으면 두고 틀리면 고쳐주세요.")
projects = [
    (12, "2025.07 — 2025.11", "Korea Environmental Conservation Institute (한국환경보전원)", "Researcher",
     "장항 국가습지복원사업 생태계조사 모니터링 용역", "(습지 모니터링 — 연결 산출물 있으면 기입)"),
    (11, "2024.06 — 2025.05", "Gyeonggi Research Institute (경기연구원)", "Researcher",
     "경기도 비오톱 유형을 고려한 탄소 저장·흡수량 분석 및 정보구축 연구",
     "제안: 산림 층위지수/탄소 관련 논문·발표 (예: KOSERT 2025 춘계 층위구조-탄소 발표)"),
    (10, "2023.04 — 2027.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "생태계서비스를 고려한 탄소흡수원 통합관리 기술개발", "제안: 산림 탄소·GEDI 관련 논문 (진행중)"),
    (9, "2022.12 — 2023.11", "Gyeonggi Research Institute (경기연구원)", "Researcher",
     "경기도 비오톱 속성정보 분석 연구", "제안: JCK 2023 경기 비오톱 속성정보 발표"),
    (8, "2022.04 — 2026.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "탄소축적 묵논습지 조성 복원 관리기술",
     "제안: Jekal et al.(2026) KOSERT(Salix 탄소) · Kim et al.(2026) AEE(묵논 곤충다양성) · "
     "특허 등록 2025(드론라이다 습지식생구조) · 특허 출원(묵논습지 육화판별)"),
    (7, "2021.09 — 2022.12", "Gyeonggi Research Institute (경기연구원)", "Researcher",
     "과천시 도시생태현황지도 작성 연구 용역",
     "제안: Lee & Song(2026) KOSERT(과천·의왕 GEDI-ALS) · KSEE 2024 과천 층위 면적 발표"),
    (6, "2021.04 — 2023.12", "Korean Ministry of Environment (환경부)", "Project Manager (책임참여)",
     "외래생물 맞춤형 실시간 웹 기반 위치정보 추적 시스템 개발",
     "제안: 야생동물 위치추적 결측점 예측 특허(2023 등록) · 최적포획범위 특허(출원)"),
    (5, "2021.04 — 2023.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "IT·ET·BT 융합 유입 외래생물 분포·확산 탐지모델 개발 연구", "(연결 산출물 있으면 기입)"),
    (4, "2020.04 — 2022.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "도시생태계 현안대응을 위한 다중기반 그린인프라 기술개발",
     "제안: UAV 야생동물 탐지 논문(Lee 2021, Remote Sensing) · 드론열화상 야생동물 특허(2022)"),
    (3, "2020.01 — 2020.08", "The Seoul Institute (서울연구원)", "Researcher",
     "공간유형별 기후환경 모니터링 및 시뮬레이션 분석",
     "제안: Jang et al.(2026) KOSERT(LST 비교) · 옥외 열쾌적성 특허(2022) · Envi-met 관련"),
    (2, "2019.04 — 2022.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "도시 생물종 맞춤형 서식환경 관리 기법 개발", "(연결 산출물 있으면 기입)"),
    (1, "2018.07 — 2020.12", "Korean Ministry of Environment (환경부)", "Researcher",
     "훼손 유형별 진단평가 체계 및 생태복원 모델 개발", "(연결 산출물 있으면 기입)"),
]
for no, period, agency, role, title_ko, outputs in projects:
    p = para(12, 0); run(p, f"[{no}] {title_ko}", 11, bold=True, color=ACCENT); bottom_border(p)
    p = para(2, 0)
    run(p, "기간: ", 10, bold=True); run(p, period + "   ", 10, color=GREY)
    run(p, "기관: ", 10, bold=True); run(p, agency + "   ", 10, color=GREY)
    run(p, "역할: ", 10, bold=True); run(p, role, 10, color=GREY)
    field("① 연구비", "(전체 또는 본인 분담 / 공개 가능 범위)")
    field("② 과제책임자(PI)", "(이름·소속)")
    p = para(2, 0); run(p, "  ③ 본인이 한 일 (분석방법·데이터·산출물 중심, 2~4개):", 10, bold=True, color=BLANK)
    for _ in range(3):
        pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(0.4)
        run(pb, "•  ________________________________________________________", 10, color=BLANK)
    p = para(2, 0); run(p, "  산출물(제안): ", 10, bold=True); run(p, outputs, 9.5, italic=True, color=GREY)
    p = para(0, 0); run(p, "  → 확인/수정: ", 10, bold=True, color=BLANK)
    run(p, "______________________________________________", 10, color=BLANK)

# ============================= PART 4 =============================
section("Part 4 — 수상 · 명예 (Awards & Honors)",
        "장학금 4건(현대차정몽구·BK21·일주·경북대)은 이미 CV Honors에 있음. 여기엔 그 외.")
subhead("◦ 학술 수상 (best paper/poster/presentation, 공모전, 학업우수상 등)")
blank_block([("연도", ""), ("수상명", ""), ("수여기관", ""),
             ("내용", "(예: Best Poster, 우수발표, 1st place)")], count=4)
subhead("◦ 학회 참가 지원금 (travel grant / 발표 지원)")
blank_block([("연도", ""), ("명칭", ""), ("지원기관", "")], count=2)

# ============================= PART 5 =============================
section("Part 5 — 학계 기여 (Service & Professional)",
        "두 예시 CV의 Service/Reviewer/Membership에 해당. 현재 CV엔 없는 섹션.")
subhead("◦ 저널 동료심사 (peer review)")
blank_block([("저널명", ""), ("리뷰 건수", ""), ("기간", "(예: 2024–현재)")], count=3)
subhead("◦ 학회 초록·세션 리뷰 / 운영 역할")
field("내용", "(있으면)")
subhead("◦ 소속 학회 회원 (AGU / ESA / KOSERT / 한국조경학회 / 한국환경생태학회 등)")
blank_block([("학회명", ""), ("회원 구분·기간", "(예: Member, 2022–현재)")], count=4)
subhead("◦ 연구실 · 학과 봉사 역할")
field("내용", "(예: 세미나 운영, 신입생 멘토, 장비 관리 등)")

# ============================= PART 6 =============================
section("Part 6 — 교육 · 멘토링 (Teaching & Mentoring detail)",
        "정규강의·멘토링은 상세화하면 강점. 온라인 4강·특강 20+는 이미 요약돼 있음.")
subhead("◦ 인천대학교 정규강의 — '공간정보의 이해와 활용' (3학점, 2024·2025)")
field("본인 역할", "(책임교수 / 강사 등 — instructor of record 여부)")
field("수강 인원", "")
field("담당 내용", "(예: 강의설계·강의·과제·평가)")
subhead("◦ 지도 · 멘토링한 후배 / 학부생")
blank_block([("이름", ""), ("수준", "(학부/석사/인턴)"),
             ("기간", ""), ("내용", "(예: LiDAR 분석 지도, 공저 논문)")], count=3)
subhead("◦ 초청 발표(invited) 구분")
p = para(2, 1); p.paragraph_format.left_indent = Inches(0.15)
run(p, "아래 중 '초청'으로 발표한 것에 표시(있으면 추가):", 10, color=BLANK)
for talk in [
    "[ ]  NEF (Neocity Empowerment Forum), Orlando 2024 — 'Remotely sensed Smartcity ecological value maps'",
    "[ ]  Asia Week, Fukuoka 2024 — 'Invasive species monitoring development and its application'",
    "[ ]  기타: " + LINE,
]:
    pb = para(1, 0); pb.paragraph_format.left_indent = Inches(0.35)
    run(pb, talk, 9.5, color=BLANK)

out = "CV_worksheet.docx"
doc.save(out)
print("saved:", out)
